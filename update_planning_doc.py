#!/usr/bin/env python3
"""
更新規劃文件：將所有已實作功能寫入「邊坡災害通報與資訊整合管理系統_規劃文件.docx」
"""

import copy
import os
from datetime import datetime

from docx import Document
from docx.shared import Pt, Emu, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

SRC = os.path.join(
    os.path.dirname(__file__),
    "邊坡災害通報與資訊整合管理系統_規劃文件.docx",
)
DST = SRC  # overwrite in-place

FONT_NAME = "微軟正黑體"
FONT_SIZE_NORMAL = Pt(12)  # 152400 EMU = 12pt
FONT_SIZE_H1 = Pt(15)     # 190500 EMU
FONT_SIZE_H2 = Pt(14)     # 177800 EMU
FONT_SIZE_H3 = Pt(13)     # 165100 EMU


# ──────────────────────────────────────────────
# helpers
# ──────────────────────────────────────────────

def set_run(run, text, font_name=FONT_NAME, font_size=FONT_SIZE_NORMAL, bold=False):
    """Apply standard formatting to a run."""
    run.text = text
    run.font.name = font_name
    run.font.size = font_size
    run.font.bold = bold
    # For CJK font
    rpr = run._element.get_or_add_rPr()
    rFonts = rpr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), font_name)


def make_paragraph(doc, text, style="Normal", font_size=FONT_SIZE_NORMAL, bold=False):
    """Create a new paragraph with standard formatting."""
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    set_run(run, text, font_size=font_size, bold=bold)
    return p


def find_paragraph_index(doc, text_contains):
    """Find the index of the first paragraph containing given text."""
    for i, p in enumerate(doc.paragraphs):
        if text_contains in p.text:
            return i
    return -1


def find_paragraph_indices(doc, text_contains):
    """Find all indices of paragraphs containing given text."""
    return [i for i, p in enumerate(doc.paragraphs) if text_contains in p.text]


def replace_paragraph_text(para, new_text, font_size=FONT_SIZE_NORMAL, bold=False):
    """Replace all text in a paragraph, preserving the paragraph element."""
    for run in para.runs:
        run.text = ""
    if para.runs:
        set_run(para.runs[0], new_text, font_size=font_size, bold=bold)
    else:
        run = para.add_run(new_text)
        set_run(run, new_text, font_size=font_size, bold=bold)


def insert_paragraph_after(doc, ref_para, text, style="Normal", font_size=FONT_SIZE_NORMAL, bold=False):
    """Insert a new paragraph after a reference paragraph."""
    new_p = doc.add_paragraph(style=style)
    run = new_p.add_run(text)
    set_run(run, text, font_size=font_size, bold=bold)
    # Move the new element to be right after ref_para
    ref_para._element.addnext(new_p._element)
    return new_p


def insert_paragraphs_after(doc, ref_para, lines, style="Normal", font_size=FONT_SIZE_NORMAL, bold=False):
    """Insert multiple paragraphs after a reference paragraph, in order."""
    last = ref_para
    for line in lines:
        last = insert_paragraph_after(doc, last, line, style=style, font_size=font_size, bold=bold)
    return last


def delete_paragraphs_between(doc, start_idx, end_idx):
    """Delete paragraphs from start_idx to end_idx (exclusive)."""
    # Must delete from bottom up to avoid index shifting
    body = doc.element.body
    to_remove = []
    for i in range(start_idx, end_idx):
        to_remove.append(doc.paragraphs[i]._element)
    for elem in to_remove:
        body.remove(elem)


def add_table_row(table, cells_text):
    """Add a row to a table with given cell texts."""
    row = table.add_row()
    for i, text in enumerate(cells_text):
        if i < len(row.cells):
            cell = row.cells[i]
            if cell.paragraphs:
                p = cell.paragraphs[0]
                if p.runs:
                    set_run(p.runs[0], text)
                else:
                    run = p.add_run(text)
                    set_run(run, text)
            else:
                p = cell.add_paragraph()
                run = p.add_run(text)
                set_run(run, text)


# ──────────────────────────────────────────────
# Section update functions
# ──────────────────────────────────────────────

def update_date(doc):
    """Update document date to today."""
    today = datetime.now().strftime("%Y年%m月%d日")
    idx = find_paragraph_index(doc, "日期：")
    if idx >= 0:
        replace_paragraph_text(doc.paragraphs[idx], f"日期：{today}")


def update_system_name(doc):
    """Update system name to reflect current WebGIS title."""
    # The system name stays the same: 邊坡災害通報與資訊整合管理系統
    # But WebGIS title is "公路即時災害通報管理系統(測試)" — this is a deployment name, not system name
    pass


def update_section_1_3(doc):
    """Update 1.3.1 Phase 1 scope to include all actually implemented features."""
    idx = find_paragraph_index(doc, "1.3.1 Phase 1")
    if idx < 0:
        return

    # Find the next heading to know our boundary
    end_idx = find_paragraph_index(doc, "1.3.2 Phase 2")
    if end_idx < 0:
        return

    # Replace content between 1.3.1 heading and 1.3.2 heading
    new_lines = [
        "- LINE Bot 通報流程：工務段選擇、路線選擇、省道里程輸入、里程轉經緯度與地圖預覽（含微調座標機制）、地質資訊確認暫停步驟、破壞模式與致災原因選擇、工程名稱填寫、照片上傳與標註、工址環境調查、初估經費、危害程度評估、確認送出。",
        "- LINE Bot 管理功能：審核待辦（動態查詢 pending_review 案件）、核准/退回/結案操作、案件刪除與 LINE 同步通知。",
        "- LINE Bot 輔助功能：案件查詢（依狀態篩選）、統計概覽、個人檔案管理。",
        "- FastAPI 後端：LINE webhook 處理、案件生命週期管理（含刪除與稽核紀錄）、Vendor API（GET/DELETE）、Word 報表自動產生、靜態地圖產製。",
        "- WebGIS 管理平台：案件點位地圖展示、工務段與狀態篩選、聚合顯示（白色氣泡）、案件詳情面板（填報摘要、照片預覽、Word 下載）、案件刪除功能、多圖層切換（含正射影像、地質敏感區等 13 圖層）。",
        "- WebGIS 統計儀表板（stats.html）：案件數量統計、工務段分布、狀態分布等視覺化圖表。",
        "- Word 報表自動產生：依「公路災害工程內容概述表」空白範本自動填入案件資料、位置簡圖（靜態地圖）、現場照片（7cm×5.25cm 連續貼入）、座標資訊，檔名以西元年月日時分命名。",
        "- 資料儲存機制：以檔案式 JSON 取代資料庫，每案獨立目錄，含稽核軌跡（audit.jsonl）。",
        "- 伺服器啟動工具：自動偵測並釋放被佔用的 port 8000，避免啟動衝突。",
    ]

    # Delete existing content paragraphs between heading and next heading
    for i in range(end_idx - 1, idx, -1):
        doc.paragraphs[i]._element.getparent().remove(doc.paragraphs[i]._element)

    # Insert new content
    insert_paragraphs_after(doc, doc.paragraphs[idx], new_lines)


def update_section_6(doc):
    """Update Chapter 6 — reporting flow to reflect actual implementation.
    Key changes:
    - Step 3/4 now include map preview + coordinate adjustment
    - Add Step 4a: CONFIRM_GEO_INFO geology info pause
    - Steps reordered to match actual implementation
    """
    # Update Step 3 description
    idx = find_paragraph_index(doc, "Step 3：GPS 座標/里程牌")
    if idx >= 0:
        # Update heading
        replace_paragraph_text(doc.paragraphs[idx], "Step 3：省道里程輸入", font_size=FONT_SIZE_H2, bold=True)
        # Update content
        content_idx = idx + 1
        if content_idx < len(doc.paragraphs) and doc.paragraphs[content_idx].style.name == "Normal":
            replace_paragraph_text(
                doc.paragraphs[content_idx],
                "使用者輸入省道里程樁號（如 45k+200），系統呼叫 LRS 服務轉換為經緯度座標，並回傳 LINE LocationMessage 供使用者在地圖上預覽位置。Quick Reply 提供三選項：「確認位置」、「重新輸入里程」、「微調座標」，其中微調座標允許使用者透過 LINE 地圖直接拖拉調整精確位置。座標顯示統一取小數點下第四位。"
            )

    # Update Step 4 description
    idx = find_paragraph_index(doc, "Step 4：里程樁號確認")
    if idx >= 0:
        replace_paragraph_text(doc.paragraphs[idx], "Step 4：地質資訊確認", font_size=FONT_SIZE_H2, bold=True)
        content_idx = idx + 1
        if content_idx < len(doc.paragraphs) and doc.paragraphs[content_idx].style.name == "Normal":
            replace_paragraph_text(
                doc.paragraphs[content_idx],
                "系統根據座標自動查詢地質敏感區與國家公園資訊，將偵測結果以文字訊息呈現。此步驟設計為暫停確認機制（CONFIRM_GEO_INFO 狀態），使用者需點選「已閱讀，繼續」按鈕後才進入下一步驟，確保重要地質資訊不被忽略。"
            )

    # Update Step 5 to include project name
    idx = find_paragraph_index(doc, "Step 5：選擇破壞模式")
    if idx >= 0:
        replace_paragraph_text(doc.paragraphs[idx], "Step 5：填寫工程名稱", font_size=FONT_SIZE_H2, bold=True)
        content_idx = idx + 1
        if content_idx < len(doc.paragraphs) and doc.paragraphs[content_idx].style.name == "Normal":
            replace_paragraph_text(
                doc.paragraphs[content_idx],
                "使用者輸入工程名稱文字，此欄位為選填，可選擇「略過後補」。工程名稱將記錄於案件資料中，並作為 Word 報表表頭資訊。"
            )

    # Update Step 6
    idx = find_paragraph_index(doc, "Step 6：選擇致災原因")
    if idx >= 0:
        replace_paragraph_text(doc.paragraphs[idx], "Step 6：選擇破壞模式", font_size=FONT_SIZE_H2, bold=True)
        content_idx = idx + 1
        if content_idx < len(doc.paragraphs) and doc.paragraphs[content_idx].style.name == "Normal":
            replace_paragraph_text(
                doc.paragraphs[content_idx],
                "三大分類：護岸/擋土牆、道路邊坡、橋梁，以 Quick Reply 點選方式呈現，作為後續致災原因篩選與模型標籤。"
            )

    # Update Step 7
    idx = find_paragraph_index(doc, "Step 7：災害描述")
    if idx >= 0:
        replace_paragraph_text(doc.paragraphs[idx], "Step 7：選擇致災原因", font_size=FONT_SIZE_H2, bold=True)
        content_idx = idx + 1
        if content_idx < len(doc.paragraphs) and doc.paragraphs[content_idx].style.name == "Normal":
            replace_paragraph_text(
                doc.paragraphs[content_idx],
                "依破壞模式過濾 28 組預定義原因，以 Quick Reply 點選方式呈現，避免自由文字造成標準不一致。"
            )

    # Update Step 8
    idx = find_paragraph_index(doc, "Step 8：上傳照片")
    if idx >= 0:
        replace_paragraph_text(doc.paragraphs[idx], "Step 8：上傳照片", font_size=FONT_SIZE_H2, bold=True)
        content_idx = idx + 1
        if content_idx < len(doc.paragraphs) and doc.paragraphs[content_idx].style.name == "Normal":
            replace_paragraph_text(
                doc.paragraphs[content_idx],
                "需上傳照片，對應 P1-P10 類型，系統即時顯示目前完成張數與各類型狀態。每張照片需標明拍攝狀況，並提供「補照片再填寫」機制，可先上傳部分照片後續再補。照片以副檔名過濾，確保僅接受圖片格式。"
            )

    # Update Step 9 → photo annotation
    idx = find_paragraph_index(doc, "Step 9：照片標註")
    if idx >= 0:
        replace_paragraph_text(doc.paragraphs[idx], "Step 9：照片標註", font_size=FONT_SIZE_H2, bold=True)
        content_idx = idx + 1
        if content_idx < len(doc.paragraphs) and doc.paragraphs[content_idx].style.name == "Normal":
            replace_paragraph_text(
                doc.paragraphs[content_idx],
                "每張照片以多選標籤標註，透過 postback 累積；P2、P4 類型為必填標註。選項盡量以點選方式提供豐富內容，需自行打字部分列為選填並提供「略過後補」選項。標註內容與 Input 內 Word 模板做勾稽驗證。"
            )

    # Update Step 10 → site survey
    idx = find_paragraph_index(doc, "Step 10：工址環境調查")
    if idx >= 0:
        content_idx = idx + 1
        if content_idx < len(doc.paragraphs) and doc.paragraphs[content_idx].style.name == "Normal":
            replace_paragraph_text(
                doc.paragraphs[content_idx],
                "填寫現地檢核清單（上邊坡、下邊坡、結構物、橋梁河道、其他），以點選方式呈現各調查項目。勾選結果形成風險特徵向量，並同步填入 Word 報表之工址環境表格。"
            )

    # Update Step 11 → cost estimate
    idx = find_paragraph_index(doc, "Step 11：初估經費")
    if idx >= 0:
        content_idx = idx + 1
        if content_idx < len(doc.paragraphs) and doc.paragraphs[content_idx].style.name == "Normal":
            replace_paragraph_text(
                doc.paragraphs[content_idx],
                "Phase 1 為選填欄位，使用者可選擇要或不要填寫。可先填區間或留空，供管理端後續補登。費用資料將同步填入 Word 報表。"
            )

    # Update Step 12 to mention Word generation
    idx = find_paragraph_index(doc, "Step 12：確認送出")
    if idx >= 0:
        content_idx = idx + 1
        if content_idx < len(doc.paragraphs) and doc.paragraphs[content_idx].style.name == "Normal":
            replace_paragraph_text(
                doc.paragraphs[content_idx],
                "以 Confirm Template 顯示摘要預覽，確認後案件送出並通知決策人員。送出後系統自動產生「公路災害工程內容概述表」Word 報表（可選），檔名以西元年月日時分命名，字體統一使用標楷體並標紅色。"
            )


def update_section_10_1(doc):
    """Update 10.1 WebGIS section with actual implementation details."""
    idx = find_paragraph_index(doc, "10.1 WebGIS")
    if idx < 0:
        return

    content_idx = idx + 1
    if content_idx < len(doc.paragraphs) and doc.paragraphs[content_idx].style.name == "Normal":
        replace_paragraph_text(
            doc.paragraphs[content_idx],
            "WebGIS 管理平台（公路即時災害通報管理系統(測試)）以 Leaflet.js 建構，提供以下功能："
        )
        # Insert detailed feature list after
        features = [
            "- 案件點位地圖展示：以 marker 標示各案件位置，支援 popup 顯示案件摘要資訊。",
            "- 聚合顯示：兩個以上案件鄰近時以白色氣泡圖示聚合，點擊可展開檢視個別案件。",
            "- 篩選功能：支援依工務段（6 段）與案件狀態（待審核、處理中、已結案、已退回）進行篩選。",
            "- 案件詳情面板：點選案件可展開右側詳情面板，包含完整填報摘要、現場照片縮圖預覽（點擊可放大）、Word 報表下載連結。",
            "- 案件刪除功能：提供紅色刪除按鈕，需二次確認，刪除後自動通知所有決策人員與案件建立者（LINE 推播），並同步更新「審核待辦」清單。",
            "- 多圖層切換：5 種底圖（正射影像為預設、電子地圖含等高線、電子地圖、混合影像、OpenStreetMap）與 8 種疊加圖層（地質敏感區(山崩地滑)、地質敏感區(全區)、道路路網、鄉鎮區界、段籍圖、國土利用現況、土壤液化潛勢、電子地圖標註透明層），資料來源為國土測繪中心 WMTS。",
            "- 統計儀表板（stats.html）：獨立統計頁面，提供案件數量、工務段分布、狀態分布等視覺化圖表。",
            "- 自動刷新：定時自動重新載入案件資料，確保地圖顯示最新狀態。",
        ]
        insert_paragraphs_after(doc, doc.paragraphs[content_idx], features)


def update_section_10_2(doc):
    """Update 10.2 Vendor API to include DELETE endpoint."""
    idx = find_paragraph_index(doc, "10.2 Vendor API")
    if idx < 0:
        return

    content_idx = idx + 1
    if content_idx < len(doc.paragraphs) and doc.paragraphs[content_idx].style.name == "Normal":
        replace_paragraph_text(
            doc.paragraphs[content_idx],
            "對外資料交換採 pull model，以 API Key（X-API-Key Header）驗證。提供以下端點：GET /api/v1/cases（增量拉取案件清單，支援 since 參數與分頁）、GET /api/v1/cases/{case_id}（單案查詢含完整內容與照片縮圖 URL）、DELETE /api/v1/cases/{case_id}（刪除案件，含稽核紀錄與 LINE 同步通知）。回傳資料均採 UTF-8 JSON 格式。"
        )


def update_api_table(doc):
    """Update the API endpoint table (Table 18) to include new endpoints."""
    if len(doc.tables) < 19:
        return
    table = doc.tables[18]

    # Check if DELETE endpoint already exists
    for row in table.rows:
        if any("DELETE" in cell.text for cell in row.cells):
            return  # Already added

    # Add new endpoints
    new_rows = [
        ["DELETE", "/api/v1/cases/{case_id}", "刪除案件（含稽核紀錄與 LINE 通知）"],
        ["GET", "/api/v1/cases/{case_id}/word", "下載案件 Word 報表（公路災害工程內容概述表）"],
        ["GET", "/api/v1/cases/{case_id}/photos/{photo_id}/thumbnail", "取得照片縮圖"],
    ]
    for row_data in new_rows:
        add_table_row(table, row_data)


def add_section_10_6(doc):
    """Add new section 10.6 Word報表自動產生."""
    # Find section 10.5 EXIF
    idx = find_paragraph_index(doc, "10.5 EXIF")
    if idx < 0:
        return

    # Find the content paragraph(s) after 10.5
    content_idx = idx + 1
    while content_idx < len(doc.paragraphs) and doc.paragraphs[content_idx].style.name == "Normal" and doc.paragraphs[content_idx].text.strip():
        content_idx += 1

    # Find the last non-empty paragraph in 10.5
    last_para = doc.paragraphs[content_idx - 1] if content_idx > idx + 1 else doc.paragraphs[idx]

    # Insert 10.6 heading and content
    h = insert_paragraph_after(doc, last_para, "10.6 Word 報表自動產生", style="Heading 2", font_size=FONT_SIZE_H2, bold=True)
    lines = [
        "系統可依據案件資料自動產生「公路災害工程內容概述表」Word 報表，以 Input 目錄下的空白 .docx 範本為基礎填入以下內容：",
        "- 表頭資訊：工程名稱、工務段、路線、里程樁號、經緯度座標（小數點下第四位）。",
        "- 位置簡圖：以 CARTO tile server 為主（OSM 為備援）產製靜態地圖截圖，標示案件點位。",
        "- 現場照片：依 P1-P10 順序貼入，尺寸固定為寬 7cm × 高 5.25cm，照片緊連貼入無標題。座標文字以段落形式附於照片下方。",
        "- 災害分析：破壞模式、致災原因、災害描述。",
        "- 工址環境調查：依上邊坡、下邊坡、結構物、橋梁河道、其他分類以勾選方式（■）填入。",
        "- 初估費用：依費用項目明細填入。",
        "- 字體統一使用標楷體並標紅色，勾選欄位以 ■ 表示（不保留 ☑ 框）。",
        "- 檔名格式：以西元年月日時分命名（如 20260301_1430.docx）。",
    ]
    insert_paragraphs_after(doc, h, lines)

    # Insert 10.7 for server management
    # Find the paragraph we just inserted (last line)
    # We need to re-find since we inserted
    idx_106 = find_paragraph_index(doc, "10.6 Word 報表自動產生")
    if idx_106 < 0:
        return
    # Go past all 10.6 content
    i = idx_106 + 1
    while i < len(doc.paragraphs) and doc.paragraphs[i].style.name == "Normal" and doc.paragraphs[i].text.strip():
        i += 1
    last_106 = doc.paragraphs[i - 1] if i > idx_106 + 1 else doc.paragraphs[idx_106]

    h2 = insert_paragraph_after(doc, last_106, "10.7 伺服器啟動管理", style="Heading 2", font_size=FONT_SIZE_H2, bold=True)
    lines2 = [
        "start_server.py 啟動腳本提供自動偵測 port 8000 佔用並釋放的功能（_kill_port），避免重複啟動時因埠號衝突導致服務無法啟動。啟動過程中的 LOG 均記錄於終端輸出，便於問題排查。",
    ]
    insert_paragraphs_after(doc, h2, lines2)


def update_section_8(doc):
    """Update Chapter 8 to mention case deletion in business track."""
    idx = find_paragraph_index(doc, "8.2 業務軌（人工）")
    if idx < 0:
        return
    content_idx = idx + 1
    if content_idx < len(doc.paragraphs) and doc.paragraphs[content_idx].style.name == "Normal":
        replace_paragraph_text(
            doc.paragraphs[content_idx],
            "業務軌狀態由決策人員操作：pending_review → in_progress → closed；若資料不足可 returned 回使用者補件，補件完成後回到 pending_review。此外，決策人員可透過 WebGIS 刪除案件，刪除前記錄完整稽核紀錄（audit log），並透過 LINE 通知所有決策人員與案件建立者。"
        )

    # Update audit trail section
    idx = find_paragraph_index(doc, "8.5 Audit Trail")
    if idx < 0:
        return
    content_idx = idx + 1
    if content_idx < len(doc.paragraphs) and doc.paragraphs[content_idx].style.name == "Normal":
        replace_paragraph_text(
            doc.paragraphs[content_idx],
            "audit.jsonl 採一行一事件格式，欄位含 timestamp、actor、action、target、before、after、source_ip、note。此格式可直接串接 SIEM 與流程分析。"
        )
    content_idx += 1
    if content_idx < len(doc.paragraphs) and doc.paragraphs[content_idx].style.name == "Normal":
        replace_paragraph_text(
            doc.paragraphs[content_idx],
            "範例事件：manager 變更 business_track 為 in_progress、user 補傳照片、系統完成 milepost_resolved、webgis_admin 刪除案件（含刪除原因與案件快照）。"
        )


def update_tech_stack_table(doc):
    """Update the technology stack table (Table 2) to add new deps."""
    if len(doc.tables) < 3:
        return
    table = doc.tables[2]

    # Check if already updated
    for row in table.rows:
        if any("python-docx" in cell.text for cell in row.cells):
            return

    new_rows = [
        ["報表產生", "python-docx", "Word 報表自動產生（填入案件資料與照片）"],
        ["靜態地圖", "staticmap + httpx", "位置簡圖產製（CARTO/OSM 底圖）"],
        ["地理服務", "行政區反查、國家公園偵測", "座標自動查詢所屬行政區與國家公園"],
    ]
    for row_data in new_rows:
        add_table_row(table, row_data)


def update_state_transition_table(doc):
    """Update state transition table (Table 8) to add deleted state."""
    if len(doc.tables) < 9:
        return
    table = doc.tables[8]

    # Check if already updated
    for row in table.rows:
        if any("deleted" in cell.text for cell in row.cells):
            return

    add_table_row(table, ["業務軌", "any", "deleted", "決策人員（WebGIS 刪除）"])


# ──────────────────────────────────────────────
# Main
# ──────────────────────────────────────────────

def main():
    print(f"[INFO] Loading: {SRC}")
    doc = Document(SRC)

    print("[1/10] Updating date...")
    update_date(doc)

    print("[2/10] Updating 1.3.1 Phase 1 scope...")
    update_section_1_3(doc)

    print("[3/10] Updating Chapter 6 reporting flow...")
    update_section_6(doc)

    print("[4/10] Updating Chapter 8 case management...")
    update_section_8(doc)

    print("[5/10] Updating 10.1 WebGIS...")
    update_section_10_1(doc)

    print("[6/10] Updating 10.2 Vendor API...")
    update_section_10_2(doc)

    print("[7/10] Adding 10.6 Word report + 10.7 Server management...")
    add_section_10_6(doc)

    print("[8/10] Updating tech stack table...")
    update_tech_stack_table(doc)

    print("[9/10] Updating state transition table...")
    update_state_transition_table(doc)

    print("[10/10] Updating API endpoint table...")
    update_api_table(doc)

    print(f"[INFO] Saving: {DST}")
    doc.save(DST)
    print("[DONE] Planning document updated successfully.")


if __name__ == "__main__":
    main()
