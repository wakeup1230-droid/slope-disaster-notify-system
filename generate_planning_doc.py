# pyright: reportUnknownVariableType=false, reportUnknownParameterType=false, reportMissingParameterType=false, reportUnknownMemberType=false, reportUnknownArgumentType=false, reportGeneralTypeIssues=false, reportMissingTypeArgument=false

from __future__ import annotations

import json
from datetime import datetime
from pathlib import Path
from typing import cast

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


OUTPUT_PATH = Path(
    r"C:\Users\Sad Kevin\Desktop\邊坡災害通報與資訊整合管理系統\邊坡災害通報與資訊整合管理系統_規劃文件.docx"
)

PHOTO_TAGS_PATH = Path(__file__).resolve().parent / "app" / "data" / "photo_tags.json"

CATEGORY_LAYOUT = [
    ("common", "7.1 共用照片類型（common）"),
    ("optional", "7.2 選用照片類型（optional）"),
    ("revetment_retaining", "7.3 護岸/擋土牆專用照片類型（revetment_retaining）"),
    ("road_slope", "7.4 道路邊坡專用照片類型（road_slope）"),
    ("bridge", "7.5 橋梁專用照片類型（bridge）"),
]


def set_run_font(run, font_name: str = "微軟正黑體", size: int = 12, bold: bool = False) -> None:
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    run.font.size = Pt(size)
    run.font.bold = bold


def add_paragraph_with_font(
    doc: Document,
    text: str,
    style: str | None = None,
    align: WD_PARAGRAPH_ALIGNMENT | None = None,
    size: int = 12,
    bold: bool = False,
) -> None:
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)


def add_heading(doc: Document, text: str, level: int) -> None:
    heading = doc.add_heading(level=level)
    run = heading.add_run(text)
    set_run_font(run, size=16 - (level * 1), bold=True)


def add_table_grid(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        p = hdr_cells[idx].paragraphs[0]
        run = p.add_run(header)
        set_run_font(run, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            p = cells[idx].paragraphs[0]
            run = p.add_run(value)
            set_run_font(run)


def add_toc_placeholder(doc: Document) -> None:
    add_heading(doc, "目錄", 1)
    add_paragraph_with_font(
        doc,
        "以下為目錄欄位，請於 Word 開啟後按 F9 更新，以顯示完整章節頁碼。",
    )

    p = doc.add_paragraph()
    # Each fldChar / instrText / t must live inside its own w:r (run).
    # Word rejects bare elements appended directly to w:p.
    run_begin = OxmlElement("w:r")
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run_begin.append(fld_begin)

    run_instr = OxmlElement("w:r")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    run_instr.append(instr)

    run_sep = OxmlElement("w:r")
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    run_sep.append(fld_sep)

    run_text = OxmlElement("w:r")
    fld_text = OxmlElement("w:t")
    fld_text.text = "更新欄位後顯示目錄"
    run_text.append(fld_text)

    run_end = OxmlElement("w:r")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run_end.append(fld_end)

    p._p.append(run_begin)
    p._p.append(run_instr)
    p._p.append(run_sep)
    p._p.append(run_text)
    p._p.append(run_end)


def add_page_break(doc: Document) -> None:
    doc.add_page_break()


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "微軟正黑體"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "微軟正黑體")
    normal.font.size = Pt(12)

    for style_name, size in [("Heading 1", 18), ("Heading 2", 15), ("Heading 3", 13)]:
        style = styles[style_name]
        style.font.name = "微軟正黑體"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "微軟正黑體")
        style.font.size = Pt(size)
        style.font.bold = True


def add_cover_page(doc: Document) -> None:
    add_paragraph_with_font(
        doc,
        "邊坡災害通報與資訊整合管理系統",
        align=WD_PARAGRAPH_ALIGNMENT.CENTER,
        size=28,
        bold=True,
    )
    add_paragraph_with_font(
        doc,
        "規劃文件",
        align=WD_PARAGRAPH_ALIGNMENT.CENTER,
        size=20,
        bold=True,
    )
    add_paragraph_with_font(doc, "", align=WD_PARAGRAPH_ALIGNMENT.CENTER)
    add_paragraph_with_font(doc, "", align=WD_PARAGRAPH_ALIGNMENT.CENTER)
    add_paragraph_with_font(
        doc,
        "機關：交通部公路局北區養護工程分局",
        align=WD_PARAGRAPH_ALIGNMENT.CENTER,
        size=14,
    )
    add_paragraph_with_font(
        doc,
        f"日期：{datetime.now().strftime('%Y年%m月%d日')}",
        align=WD_PARAGRAPH_ALIGNMENT.CENTER,
        size=14,
    )
    add_paragraph_with_font(
        doc,
        "版本：v1.1.0",
        align=WD_PARAGRAPH_ALIGNMENT.CENTER,
        size=14,
    )


def chapter_1(doc: Document) -> None:
    add_heading(doc, "第1章 專案概述", 1)
    add_heading(doc, "1.1 系統名稱與建置背景", 2)
    add_paragraph_with_font(
        doc,
        "系統名稱為「邊坡災害通報與資訊整合管理系統」，目的是建立一套可由第一線巡查人員快速回報、主管即時掌握、後續可擴充 AI 分析能力的災害資訊平台。",
    )
    add_paragraph_with_font(
        doc,
        "本系統定位為北區養護工程分局之跨工務段協作工具，整合 LINE Bot 回報入口、案件資料歸檔、位置判讀、照片標註與決策審核流程，以縮短災情通報到處置決策的時間差。",
    )

    add_heading(doc, "1.2 專案目標", 2)
    goals = [
        "建構公路局北區養護工程分局邊坡災害通報 LINE Bot 系統，提供低門檻、低學習成本的通報流程。",
        "建立統一案件資料結構，確保第一線回報品質、可追溯性與後續介接能力。",
        "提供決策人員案件審核與查詢介面，支援跨工務段比對與優先序判斷。",
        "在 Phase 1 即完成 AI 可用資料設計，為 Phase 2 推論模型建立高品質標註基礎。",
    ]
    for item in goals:
        add_paragraph_with_font(doc, f"• {item}")

    add_heading(doc, "1.3 階段範圍", 2)
    add_heading(doc, "1.3.1 Phase 1（Vibe Coding）範圍", 3)
    p1 = [
        "LINE Bot 通報流程：工務段選擇、路線選擇、省道里程輸入、里程轉經緯度與地圖預覽（含微調座標機制）、地質資訊確認暫停步驟、破壞模式與致災原因選擇、工程名稱填寫、照片上傳與標註、工址環境調查、初估經費、危害程度評估、確認送出。",
        "LINE Bot 管理功能：審核待辦（動態查詢 pending_review 案件）、核准/退回/結案操作、案件刪除與 LINE 同步通知。",
        "LINE Bot 輔助功能：案件查詢（依狀態篩選）、統計概覽、個人檔案管理。",
        "FastAPI 後端：LINE webhook 處理、案件生命週期管理（含刪除與稽核紀錄）、Vendor API（GET/DELETE）、Word 報表自動產生、靜態地圖產製。",
        "WebGIS 管理平台：案件點位地圖展示、工務段與狀態篩選、聚合顯示（白色氣泡）、案件詳情面板（填報摘要、照片預覽、Word 下載）、案件刪除功能、多圖層切換（含正射影像、地質敏感區等 13 圖層）。",
        "WebGIS 統計儀表板（stats.html）：案件數量統計、工務段分布、狀態分布等視覺化圖表。",
        "Word 報表自動產生：依「公路災害工程內容概述表」空白範本自動填入案件資料、位置簡圖（靜態地圖）、現場照片（7cm×5.25cm 連續貼入）、座標資訊，檔名以西元年月日時分命名。",
        "資料儲存機制：以檔案式 JSON 取代資料庫，每案獨立目錄，含稽核軌跡（audit.jsonl）。",
        "伺服器啟動工具：自動偵測並釋放被佔用的 port 8000，避免啟動衝突。",
    ]
    for item in p1:
        add_paragraph_with_font(doc, f"- {item}")

    add_heading(doc, "1.3.2 Phase 2 範圍預覽", 3)
    p2 = [
        "AI/LLM 推論整合：多模態資料（照片、文字、座標、里程）聯合分析。",
        "自動分類：根據照片與標註推斷破壞模式與致災原因。",
        "嚴重度預測：生成標準化分級建議，輔助主管排序處置優先序。",
        "報表自動化：產生摘要、說明文字與案件趨勢分析。",
    ]
    for item in p2:
        add_paragraph_with_font(doc, f"- {item}")

    add_heading(doc, "1.4 成功指標", 2)
    indicators = [
        "通報流程平均完成時間低於 8 分鐘（含照片上傳）。",
        "案件資料欄位完整率達 95% 以上。",
        "里程樁號自動判讀成功率達 85% 以上（可接受人工修正）。",
        "決策人員可在 2 分鐘內完成案件初審並回饋狀態。",
        "所有案件留下可追溯 audit trail，符合稽核需求。",
    ]
    for item in indicators:
        add_paragraph_with_font(doc, f"- {item}")


def chapter_2(doc: Document) -> None:
    add_heading(doc, "第2章 系統架構", 1)
    add_heading(doc, "2.1 單體 FastAPI 架構", 2)
    add_paragraph_with_font(
        doc,
        "Phase 1 採單一 FastAPI 應用程式架構，透過明確目錄分層維持可維護性；應用程式入口集中在 app.py，業務邏輯由 routers、services、models 與 core 模組分工。",
    )
    add_table_grid(
        doc,
        ["模組", "說明", "責任範圍"],
        [
            ["routers/", "API 路由層", "接收請求、驗證輸入、回傳標準回應"],
            ["services/", "服務層", "執行流程邏輯、資料計算、狀態轉換"],
            ["models/", "資料模型", "Pydantic schema、列舉值、型別約束"],
            ["core/", "核心共用", "設定載入、日誌、權限、工具函式"],
        ],
    )

    add_heading(doc, "2.2 檔案式資料儲存", 2)
    add_paragraph_with_font(
        doc,
        "Phase 1 不使用資料庫，採 file-based JSON 儲存，以便快速部署與離線備援。每個案件以獨立目錄管理，避免單點檔案過大並利於稽核。",
    )
    add_paragraph_with_font(doc, "案件目錄命名規範：case_YYYYMMDD_NNNN")
    add_table_grid(
        doc,
        ["檔案/目錄", "用途", "內容摘要"],
        [
            ["case.json", "主資料", "案件欄位、通報內容、雙軌狀態、派工資訊"],
            ["audit.jsonl", "稽核軌跡", "逐筆事件紀錄，含時間、操作者、前後值"],
            ["evidence/", "證據原檔", "原始照片與附件，依 SHA-256 命名"],
            ["derived/", "衍生資料", "OCR 結果、LRS（Linear Referencing System，線性參考系統）判讀結果"],
            ["thumbnails/", "縮圖快取", "查詢列表與審核畫面所需縮圖"],
        ],
    )

    add_heading(doc, "2.3 LINE Messaging API 整合", 2)
    add_paragraph_with_font(
        doc,
        "LINE webhook 由 /line/webhook 端點接收，使用 line-bot-sdk 驗簽與事件解析。事件包含文字、位置、圖片、postback、follow/unfollow，依角色進入不同狀態機。",
    )
    add_paragraph_with_font(
        doc,
        "為降低對話中斷風險，後端於每一步儲存 session_state，當使用者中途離開後可從最近步驟繼續；多張照片上傳則採批次累積並即時回饋目前張數。",
    )

    add_heading(doc, "2.4 證據檔案內容定址（Content Addressed）", 2)
    add_paragraph_with_font(
        doc,
        "證據檔案以 SHA-256 當作內容識別碼，檔名格式為 sha256.<副檔名>，可自然去重並降低重複儲存成本。案件中僅存參照關係與 metadata。",
    )
    add_paragraph_with_font(
        doc,
        "每筆 evidence metadata 至少包含 uploader、captured_at、mime_type、bytes、sha256、source（line_upload/manual）與鏈結狀態。",
    )

    add_heading(doc, "2.5 技術堆疊", 2)
    add_table_grid(
        doc,
        ["分類", "技術", "用途"],
        [
            ["語言與執行環境", "Python 3.10+", "主程式語言與依賴管理"],
            ["API 框架", "FastAPI + uvicorn", "REST API 與非同步 I/O"],
            ["通訊整合", "line-bot-sdk", "LINE Messaging API webhook 與訊息回覆"],
            ["影像處理", "Pillow", "縮圖、基本影像前處理"],
            ["座標轉換", "pyproj", "座標系轉換與距離計算前處理"],
            ["資料分析", "pandas", "彙整查詢與匯出報表"],
            ["外部呼叫", "httpx", "Vendor API 與內部服務整合"],
            ["PDF 解析", "PyMuPDF、pdfplumber", "工程文件文字/表格擷取"],
            ["OCR", "pytesseract", "里程牌與關鍵文字辨識"],
            ["地圖前端", "Leaflet.js", "WebGIS 標記、聚合與篩選"],
            ["報表產生", "python-docx", "Word 報表自動產生（填入案件資料與照片）"],
            ["靜態地圖", "staticmap + httpx", "位置簡圖產製（CARTO/OSM 底圖）"],
            ["地理服務", "行政區反查、國家公園偵測", "座標自動查詢所屬行政區與國家公園"],
        ],
    )


def chapter_3(doc: Document) -> None:
    add_heading(doc, "第3章 使用者角色與權限", 1)
    add_heading(doc, "3.1 角色設計原則", 2)
    add_paragraph_with_font(
        doc,
        "系統僅定義兩種角色：使用者人員與決策人員。為避免權責模糊，本期不設「承辦人」角色，所有流程在這兩種角色內完成。",
    )
    add_table_grid(
        doc,
        ["角色", "主要任務", "核心權限"],
        [
            ["使用者人員", "災情通報、補件、查詢個人案件", "建立案件、上傳證據、查看個人案件狀態"],
            ["決策人員", "審核、派工、跨段監看、結案", "查詢全段案件、更新業務狀態、退回補件"],
        ],
    )

    add_heading(doc, "3.2 註冊與審核流程", 2)
    flow = [
        "使用者透過 LINE Follow 觸發註冊流程，填寫真實姓名與工務段別。",
        "系統建立待審核帳號，status 設為 pending。",
        "決策人員收到審核通知，核准後 status 變更為 active。",
        "若資訊不完整可退回，使用者補件後再送審。",
        "核准事件記錄 approved_at、approved_by 以供追溯。",
    ]
    for item in flow:
        add_paragraph_with_font(doc, f"- {item}")

    add_heading(doc, "3.3 Bootstrap Admin", 2)
    add_paragraph_with_font(
        doc,
        "系統首次部署由設定檔注入 bootstrap admin（決策人員）帳號，僅限啟用初期使用。啟用後應建立正式管理帳號並停用 bootstrap 機制以符合資安要求。",
    )

    add_heading(doc, "3.4 Rich Menu 與權限綁定", 2)
    add_paragraph_with_font(
        doc,
        "角色切換後自動綁定對應 Rich Menu。使用者人員看見通報與個案查詢功能；決策人員看見審核、統計、跨段監看與設定功能。",
    )

    add_heading(doc, "3.5 User JSON 結構", 2)
    add_table_grid(
        doc,
        ["欄位", "型別", "說明"],
        [
            ["user_id", "string", "LINE user id，系統主鍵"],
            ["display_name", "string", "LINE 顯示名稱"],
            ["real_name", "string", "真實姓名（註冊填寫）"],
            ["district_id", "string", "工務段識別碼"],
            ["district_name", "string", "工務段名稱"],
            ["role", "string", "user 或 manager"],
            ["status", "string", "pending/active/rejected/disabled"],
            ["registered_at", "datetime", "註冊時間"],
            ["approved_at", "datetime|null", "核准時間"],
            ["approved_by", "string|null", "核准者 user_id"],
        ],
    )


def chapter_4(doc: Document) -> None:
    add_heading(doc, "第4章 組織架構 — 六工務段", 1)
    add_paragraph_with_font(
        doc,
        "本系統以北區養護工程分局六工務段為基本管理單位，案件歸屬、查詢權限與統計維度均以工務段為主鍵。",
    )
    add_table_grid(
        doc,
        ["工務段", "ID", "管轄公路"],
        [
            ["景美工務段", "jingmei", "台2(部分), 台3, 台5, 台9, 台9甲, 台9乙"],
            ["中和工務段", "zhonghe", "台3(都會), 台9(都會), 台15(部分), 台64, 台65"],
            ["中壢工務段", "zhongli", "台1, 台4, 台15(桃園), 台31, 台66, 台61(部分)"],
            ["新竹工務段", "hsinchu", "台1(新竹), 台3(部分), 台13, 台15(濱海), 台31(部分), 台61(部分), 台68"],
            ["復興工務段", "fuxing", "台7, 台7甲, 台7乙, 台118(部分), 台4(東段), 台3(北端山區)"],
            ["基隆工務段", "keelung", "台2, 台2甲(北段), 台5(基隆-汐止), 台9(瑞芳短段), 台62, 台62甲, 台102"],
        ],
    )
    add_paragraph_with_font(
        doc,
        "工務段資料將同時用於：註冊核准範圍、路線篩選、案件查詢、統計面板分群、Vendor API 同步參數。",
    )


def chapter_5(doc: Document) -> None:
    add_heading(doc, "第5章 LINE Bot 介面設計", 1)
    add_heading(doc, "5.1 Rich Menu 版面", 2)
    add_paragraph_with_font(doc, "兩種角色均採 6 宮格按鈕，降低操作記憶負擔。")
    add_table_grid(
        doc,
        ["角色", "按鈕 1", "按鈕 2", "按鈕 3", "按鈕 4", "按鈕 5", "按鈕 6"],
        [
            ["使用者人員", "我要通報", "我的案件", "草稿箱", "補件上傳", "常見問題", "聯絡窗口"],
            ["決策人員", "待審核案件", "全段總覽", "案件搜尋", "GIS 地圖", "指標儀表板", "系統設定"],
        ],
    )

    add_heading(doc, "5.2 Flex Message 模板", 2)
    flex_points = [
        "案件卡片模板：顯示案件編號、工務段、路線、最新狀態、嚴重度與縮圖。",
        "審核摘要模板：以四區塊顯示照片摘要、位置分析、文字描述、資料品質燈號。",
        "通知模板：核准、退回、補件完成、狀態變更均用統一版型降低辨識成本。",
    ]
    for item in flex_points:
        add_paragraph_with_font(doc, f"- {item}")

    add_heading(doc, "5.3 Quick Reply 與 Confirm Template", 2)
    add_paragraph_with_font(
        doc,
        "Quick Reply 用於高頻、低複雜度選項（工務段、路線、步驟型單選）；Confirm Template 用於不可逆操作（送出案件、關閉案件、退回補件）。",
    )

    add_heading(doc, "5.4 Postback 多選累積模式", 2)
    add_paragraph_with_font(
        doc,
        "照片標註採 postback accumulation pattern：每次點選以 postback 傳入 tag，後端更新暫存集合並回傳目前已選標籤。使用者按「完成此照片」後才提交，避免誤送。",
    )
    add_paragraph_with_font(
        doc,
        "此模式可降低長表單輸入負擔，並保留可重入能力。若連線中斷，重新開啟會以 server-side state 還原已選項目。",
    )


def chapter_6(doc: Document) -> None:
    add_heading(doc, "第6章 災害通報流程（12步驟）", 1)
    steps = [
        ("Step 1：選擇工務段", "以 Quick Reply 顯示 6 工務段，選擇後建立案件草稿並寫入 district_id。"),
        ("Step 2：選擇路線", "根據工務段過濾路線清單，避免非管轄路線誤選，並寫入 route_id 與 route_name。"),
        (
            "Step 3：省道里程輸入",
            "使用者輸入省道里程樁號（如 45k+200），系統呼叫 LRS 服務轉換為經緯度座標，並回傳 LINE LocationMessage 供使用者在地圖上預覽位置。Quick Reply 提供三選項：「確認位置」、「重新輸入里程」、「微調座標」，其中微調座標允許使用者透過 LINE 地圖直接拖拉調整精確位置。座標顯示統一取小數點下第四位。",
        ),
        (
            "Step 4：地質資訊確認",
            "系統根據座標自動查詢地質敏感區與國家公園資訊，將偵測結果以文字訊息呈現。此步驟設計為暫停確認機制（CONFIRM_GEO_INFO 狀態），使用者需點選「已閱讀，繼續」按鈕後才進入下一步驟，確保重要地質資訊不被忽略。",
        ),
        (
            "Step 5：填寫工程名稱",
            "使用者輸入工程名稱文字，此欄位為選填，可選擇「略過後補」。工程名稱將記錄於案件資料中，並作為 Word 報表表頭資訊。",
        ),
        (
            "Step 6：選擇破壞模式",
            "三大分類：護岸/擋土牆、道路邊坡、橋梁，以 Quick Reply 點選方式呈現，作為後續致災原因篩選與模型標籤。",
        ),
        (
            "Step 7：選擇致災原因",
            "依破壞模式過濾 28 組預定義原因，以 Quick Reply 點選方式呈現，避免自由文字造成標準不一致。",
        ),
        (
            "Step 8：上傳照片",
            "需上傳照片，對應 P1-P10 類型，系統即時顯示目前完成張數與各類型狀態。每張照片需標明拍攝狀況，並提供「補照片再填寫」機制，可先上傳部分照片後續再補。照片以副檔名過濾，確保僅接受圖片格式。",
        ),
        (
            "Step 9：照片標註",
            "每張照片以多選標籤標註，透過 postback 累積；P2、P4 類型為必填標註。選項盡量以點選方式提供豐富內容，需自行打字部分列為選填並提供「略過後補」選項。標註內容與 Input 內 Word 模板做勾稽驗證。",
        ),
        (
            "Step 10：工址環境調查",
            "填寫現地檢核清單（上邊坡、下邊坡、結構物、橋梁河道、其他），以點選方式呈現各調查項目。勾選結果形成風險特徵向量，並同步填入 Word 報表之工址環境表格。",
        ),
        (
            "Step 11：初估經費",
            "Phase 1 為選填欄位，使用者可選擇填寫或略過。可先填區間或留空，供管理端後續補登。費用資料將同步填入 Word 報表。",
        ),
        (
            "Step 12：確認送出",
            "以 Confirm Template 顯示摘要預覽，確認後案件送出並通知決策人員。送出後系統自動產生「公路災害工程內容概述表」Word 報表（可選），檔名以西元年月日時分命名，字體統一使用標楷體並標紅色。",
        ),
    ]
    for title, content in steps:
        add_heading(doc, title, 2)
        add_paragraph_with_font(doc, content)

    add_heading(doc, "6.13 通報例外處理", 2)
    exceptions = [
        "照片不足：顯示缺少類型與張數，禁止送出。",
        "里程信心值低：系統要求二次確認並保留候選清單。",
        "網路不穩：暫存草稿與分段補傳，避免資料遺失。",
        "重複案件：依時間、距離、路線比對給出疑似重複警示。",
    ]
    for item in exceptions:
        add_paragraph_with_font(doc, f"- {item}")


def chapter_7(doc: Document) -> None:
    add_heading(doc, "第7章 照片類型與標註系統", 1)
    add_paragraph_with_font(
        doc,
        "本系統共定義 10 種照片類型（P1-P10），分布於 5 大分類（common、optional、revetment_retaining、road_slope、bridge），以下完整列出各分類之照片類型、標註分類與全部可選項目。",
    )

    with PHOTO_TAGS_PATH.open("r", encoding="utf-8") as f:
        raw_data = cast(object, json.load(f))

    if not isinstance(raw_data, dict):
        raise ValueError("photo_tags.json 格式錯誤：根層必須為物件")

    photo_tags_data: dict[str, object] = raw_data

    def as_dict_list(value: object) -> list[dict[str, object]]:
        if not isinstance(value, list):
            return []
        result: list[dict[str, object]] = []
        for item in value:
            if isinstance(item, dict):
                result.append(item)
        return result

    def as_text(value: object, default: str = "") -> str:
        if isinstance(value, str):
            return value
        if isinstance(value, (int, float)):
            return str(value)
        return default

    def as_bool(value: object) -> bool:
        return value is True

    def tag_line(tag_def: dict[str, object]) -> str:
        mode = "多選" if as_bool(tag_def.get("multi_select")) else "單選"
        tier = as_text(tag_def.get("tier"), "unknown")
        category_name = as_text(tag_def.get("category_name"), "未命名分類")
        return f"- {category_name}（{mode}，tier: {tier}）"

    def options_line(prefix: str, tags: list[dict[str, object]]) -> str:
        if not tags:
            return f"  {prefix}：無"
        labels = "、".join(as_text(tag.get("label")) for tag in tags)
        return f"  {prefix}：{labels}"

    for category_key, category_heading in CATEGORY_LAYOUT:
        add_heading(doc, category_heading, 2)
        category_block_raw = photo_tags_data.get(category_key, {})
        if not isinstance(category_block_raw, dict):
            continue

        for photo_type_key, photo_type_def_raw in category_block_raw.items():
            if not isinstance(photo_type_key, str):
                continue
            if not isinstance(photo_type_def_raw, dict):
                continue

            required = "必填" if as_bool(photo_type_def_raw.get("required")) else "選填"
            max_photos = as_text(photo_type_def_raw.get("max_photos"), "-")
            add_heading(
                doc,
                f"{photo_type_key} {as_text(photo_type_def_raw.get('name'))}（{required}，最多 {max_photos} 張）",
                3,
            )

            add_paragraph_with_font(doc, "照片標註項目 (photo_tags)", bold=True)
            photo_tag_groups = as_dict_list(photo_type_def_raw.get("photo_tags", []))
            if not photo_tag_groups:
                add_paragraph_with_font(doc, "- 無")
            for tag_group in photo_tag_groups:
                add_paragraph_with_font(doc, tag_line(tag_group))
                add_paragraph_with_font(doc, options_line("選項", as_dict_list(tag_group.get("tags", []))))
                exclusion_tags = as_dict_list(tag_group.get("exclusion_tags", []))
                if exclusion_tags:
                    add_paragraph_with_font(
                        doc,
                        options_line("排除選項", exclusion_tags),
                    )

            add_paragraph_with_font(doc, "人工研判項目 (judgment_tags)", bold=True)
            judgment_tag_groups = as_dict_list(photo_type_def_raw.get("judgment_tags", []))
            if not judgment_tag_groups:
                add_paragraph_with_font(doc, "- 無")
            for tag_group in judgment_tag_groups:
                base_line = tag_line(tag_group)
                if as_text(tag_group.get("input_type")) == "text":
                    base_line += "（自由文字輸入）"
                add_paragraph_with_font(doc, base_line)
                if as_text(tag_group.get("input_type")) == "text":
                    add_paragraph_with_font(doc, "  選項：自由文字輸入")
                else:
                    add_paragraph_with_font(doc, options_line("選項", as_dict_list(tag_group.get("tags", []))))
                exclusion_tags = as_dict_list(tag_group.get("exclusion_tags", []))
                if exclusion_tags:
                    add_paragraph_with_font(
                        doc,
                        options_line("排除選項", exclusion_tags),
                    )


def chapter_8(doc: Document) -> None:
    add_heading(doc, "第8章 案件管理與雙軌狀態系統", 1)
    add_heading(doc, "8.1 技術軌（自動）", 2)
    add_paragraph_with_font(
        doc,
        "技術軌狀態由系統程序驅動：ingested → photos_processed → milepost_resolved → complete。任一步驟失敗將保留錯誤資訊並可重試。",
    )

    add_heading(doc, "8.2 業務軌（人工）", 2)
    add_paragraph_with_font(
        doc,
        "業務軌狀態由決策人員操作：pending_review → in_progress → closed；若資料不足可 returned 回使用者補件，補件完成後回到 pending_review。此外，決策人員可透過 WebGIS 刪除案件，刪除前記錄完整稽核紀錄（audit log），並透過 LINE 通知所有決策人員與案件建立者。",
    )

    add_heading(doc, "8.3 合法轉換規則", 2)
    add_table_grid(
        doc,
        ["軌道", "起始狀態", "可轉移至", "觸發者"],
        [
            ["技術軌", "ingested", "photos_processed", "系統背景程序"],
            ["技術軌", "photos_processed", "milepost_resolved", "LRS 服務"],
            ["技術軌", "milepost_resolved", "complete", "系統背景程序"],
            ["業務軌", "pending_review", "in_progress / returned", "決策人員"],
            ["業務軌", "in_progress", "closed / returned", "決策人員"],
            ["業務軌", "returned", "pending_review", "使用者補件後系統轉換"],
            ["業務軌", "any", "deleted", "決策人員（WebGIS 刪除）"],
        ],
    )

    add_heading(doc, "8.4 Case JSON 欄位總覽", 2)
    case_fields = [
        ["case_id", "string", "案件編號，例：case_20260228_0007"],
        ["district", "object", "工務段資訊（id/name）"],
        ["route", "object", "路線、方向、里程區段"],
        ["location", "object", "經緯度、座標系、來源、信心值"],
        ["damage", "object", "破壞模式、致災原因、文字描述"],
        ["photos", "array", "P1-P10 照片與標註列表"],
        ["survey", "object", "工址環境調查清單"],
        ["cost_estimate", "object|null", "初估經費（可空）"],
        ["status", "object", "technical_track、business_track"],
        ["timeline", "object", "建立、送出、結案等時間"],
    ]
    add_table_grid(doc, ["欄位", "型別", "說明"], case_fields)

    add_heading(doc, "8.5 Audit Trail（audit.jsonl）", 2)
    add_paragraph_with_font(
        doc,
        "audit.jsonl 採一行一事件格式，欄位含 timestamp、actor、action、target、before、after、source_ip、note。此格式可直接串接 SIEM 與流程分析。",
    )
    add_paragraph_with_font(
        doc,
        "範例事件：manager 變更 business_track 為 in_progress、user 補傳照片、系統完成 milepost_resolved、webgis_admin 刪除案件（含刪除原因與案件快照）。",
    )


def chapter_9(doc: Document) -> None:
    add_heading(doc, "第9章 案件查詢與決策人員審核介面", 1)
    add_heading(doc, "9.1 使用者查詢介面", 2)
    add_paragraph_with_font(
        doc,
        "使用者人員可查詢我的案件（my cases），並依狀態篩選：pending_review、in_progress、returned、closed；每筆案件顯示最後更新時間與補件要求。",
    )

    add_heading(doc, "9.2 決策人員查詢介面", 2)
    add_paragraph_with_font(
        doc,
        "決策人員可依工務段、路線、時間區間與嚴重度搜尋，並可切換全段總覽模式掌握跨區災害分布與案件壅塞狀況。",
    )

    add_heading(doc, "9.3 審核四視圖", 2)
    views = [
        "照片視圖：瀏覽 P1-P10 與標註，快速檢查資料完整度。",
        "地圖視圖：定位案件點位、鄰近案件與熱點標示。",
        "文字視圖：閱讀災害描述、現場調查、補件說明。",
        "位置分析視圖：顯示 GPS/里程一致性與 confidence。",
    ]
    for item in views:
        add_paragraph_with_font(doc, f"- {item}")

    add_heading(doc, "9.4 資料品質指標", 2)
    add_table_grid(
        doc,
        ["指標", "判定規則", "燈號"],
        [
            ["必填照片完整度", "P1-P4 是否齊全", "綠/黃/紅"],
            ["位置可信度", "LRS confidence 門檻", "綠/黃/紅"],
            ["描述充分性", "文字長度與關鍵欄位覆蓋", "綠/黃/紅"],
            ["標註一致性", "標籤與破壞模式一致程度", "綠/黃/紅"],
        ],
    )

    add_heading(doc, "9.5 案件列表訊息格式", 2)
    add_paragraph_with_font(
        doc,
        "LINE 端採 Carousel/Flex 呈現案件列表，每張卡片包含案件編號、地點、狀態、嚴重度與快捷操作（查看、補件、審核）。",
    )


def chapter_10(doc: Document) -> None:
    add_heading(doc, "第10章 系統整合", 1)
    add_heading(doc, "10.1 WebGIS 示範頁（Leaflet.js）", 2)
    add_paragraph_with_font(
        doc,
        "WebGIS 管理平台（公路即時災害通報管理系統(測試)）以 Leaflet.js 建構，提供以下功能：",
    )
    webgis_features = [
        "案件點位地圖展示：以 marker 標示各案件位置，支援 popup 顯示案件摘要資訊。",
        "聚合顯示：兩個以上案件鄰近時以白色氣泡圖示聚合，點擊可展開檢視個別案件。",
        "篩選功能：支援依工務段（6 段）與案件狀態（待審核、處理中、已結案、已退回）進行篩選。",
        "案件詳情面板：點選案件可展開右側詳情面板，包含完整填報摘要、現場照片縮圖預覽（點擊可放大）、Word 報表下載連結。",
        "案件刪除功能：提供紅色刪除按鈕，需二次確認，刪除後自動通知所有決策人員與案件建立者（LINE 推播），並同步更新「審核待辦」清單。",
        "多圖層切換：5 種底圖（正射影像為預設、電子地圖含等高線、電子地圖、混合影像、OpenStreetMap）與 8 種疊加圖層（地質敏感區(山崩地滑)、地質敏感區(全區)、道路路網、鄉鎮區界、段籍圖、國土利用現況、土壤液化潛勢、電子地圖標註透明層），資料來源為國土測繪中心 WMTS。",
        "統計儀表板（stats.html）：獨立統計頁面，提供案件數量、工務段分布、狀態分布等視覺化圖表。",
        "自動刷新：定時自動重新載入案件資料，確保地圖顯示最新狀態。",
    ]
    for item in webgis_features:
        add_paragraph_with_font(doc, f"- {item}")

    add_heading(doc, "10.2 Vendor API（Pull Model）", 2)
    add_paragraph_with_font(
        doc,
        "對外資料交換採 pull model，以 API Key（X-API-Key Header）驗證。提供以下端點：GET /api/v1/cases（增量拉取案件清單，支援 since 參數與分頁）、GET /api/v1/cases/{case_id}（單案查詢含完整內容與照片縮圖 URL）、DELETE /api/v1/cases/{case_id}（刪除案件，含稽核紀錄與 LINE 同步通知）。回傳資料均採 UTF-8 JSON 格式。",
    )
    add_table_grid(
        doc,
        ["項目", "規格"],
        [
            ["認證方式", "X-API-Key Header"],
            ["同步模式", "依 since 參數做增量拉取"],
            ["資料格式", "JSON（UTF-8）"],
            ["錯誤處理", "HTTP status + error_code + message"],
        ],
    )

    add_heading(doc, "10.3 PDF 解析管線", 2)
    add_paragraph_with_font(
        doc,
        "PDF parser pipeline 使用 PyMuPDF 擷取版面文字、pdfplumber 解析表格、pytesseract 處理掃描頁。流程會輸出結構化 JSON 供案件比對與歷史資料補值。",
    )

    add_heading(doc, "10.4 LRS（Linear Referencing System）線性參考系統服務", 2)
    add_paragraph_with_font(
        doc,
        "LRS 採 grid-hash + geodesic distance 計算候選里程，依距離、路線方向、歷史一致性加權得出 confidence scoring；若信心不足則要求人工確認。",
    )

    add_heading(doc, "10.5 Word 報表自動產生", 2)
    add_paragraph_with_font(
        doc,
        "系統可依據案件資料自動產生「公路災害工程內容概述表」Word 報表，以 Input 目錄下的空白 .docx 範本為基礎填入以下內容：",
    )
    word_points = [
        "表頭資訊：工程名稱、工務段、路線、里程樁號、經緯度座標（小數點下第四位）。",
        "位置簡圖：以 CARTO tile server 為主（OSM 為備援）產製靜態地圖截圖，標示案件點位。",
        "現場照片：依 P1-P10 順序貼入，尺寸固定為寬 7cm × 高 5.25cm，照片緊連貼入無標題。座標文字以段落形式附於照片下方。",
        "災害分析：破壞模式、致災原因、災害描述。",
        "工址環境調查：依上邊坡、下邊坡、結構物、橋梁河道、其他分類以勾選方式（■）填入。",
        "初估費用：依費用項目明細填入。",
        "字體統一使用標楷體並標紅色，勾選欄位以 ■ 表示（不保留 ☑ 框）。",
        "檔名格式：以西元年月日時分命名（如 20260301_1430.docx）。",
    ]
    for item in word_points:
        add_paragraph_with_font(doc, f"- {item}")

    add_heading(doc, "10.6 伺服器啟動管理", 2)
    add_paragraph_with_font(
        doc,
        "start_server.py 啟動腳本提供自動偵測 port 8000 佔用並釋放的功能（_kill_port），避免重複啟動時因埠號衝突導致服務無法啟動。啟動過程中的 LOG 均記錄於終端輸出，便於問題排查。",
    )


def chapter_11(doc: Document) -> None:
    add_heading(doc, "第11章 第二階段 AI 藍圖", 1)
    add_heading(doc, "11.1 Phase 1 對 AI 的資料價值", 2)
    value_rows = [
        ["照片標註", "損壞分類 CNN 訓練資料", "提升自動辨識準確率"],
        ["結構化標籤", "多標籤分類資料集", "降低人工分類負擔"],
        ["嚴重度評分", "迴歸/序位模型目標值", "建立客觀分級建議"],
        ["GPS+里程", "空間機器學習特徵", "熱點預測與風險評分"],
        ["破壞模式+原因", "決策樹/知識圖譜", "建立可解釋規則"],
        ["Audit Trail", "流程探勘、SLA 預測", "發現流程瓶頸"],
        ["工址調查清單", "風險評估特徵向量", "支援預警模型"],
    ]
    add_table_grid(doc, ["Phase 1 資料", "Phase 2 模型用途", "預期效益"], value_rows)

    add_heading(doc, "11.2 第二階段功能規劃", 2)
    features = [
        "自動災損分類：由照片與標註推斷破壞模式。",
        "多模態嚴重度預測：整合圖片、文字、位置、環境調查。",
        "優先處置建議引擎：依風險、影響、資源可用性產生排序。",
        "空間風險熱圖：以歷史案件與地形因子生成區域風險圖。",
        "自然語言報告生成：自動輸出主管摘要與對外說明版本。",
        "通報異常偵測：辨識異常回報模式、重複通報或誤報。",
    ]
    for item in features:
        add_paragraph_with_font(doc, f"- {item}")

    add_heading(doc, "11.3 模型治理與資料倫理", 2)
    governance = [
        "模型輸出需保留可解釋特徵，避免黑箱影響決策責任。",
        "資料匿名化與最小必要原則，避免個資外洩風險。",
        "建立模型監控與再訓練週期，處理資料漂移。",
        "保留人工覆核機制，不以 AI 直接取代決策人員。",
    ]
    for item in governance:
        add_paragraph_with_font(doc, f"- {item}")


def chapter_12(doc: Document) -> None:
    add_heading(doc, "第12章 實施計畫與時程", 1)
    add_heading(doc, "12.1 15 項任務與相依關係", 2)
    tasks = [
        ["T01", "需求確認與欄位凍結", "-", "A"],
        ["T02", "LINE Channel 與 webhook 建置", "T01", "A"],
        ["T03", "使用者註冊與審核流程", "T01", "A"],
        ["T04", "通報 12 步驟狀態機", "T02,T03", "A"],
        ["T05", "照片上傳與 SHA-256 儲存", "T04", "B"],
        ["T06", "照片標註與多選累積機制", "T05", "B"],
        ["T07", "LRS 轉換服務與信心評分", "T04", "B"],
        ["T08", "Case JSON/Audit JSONL 實作", "T04", "B"],
        ["T09", "決策人員審核流程", "T08", "C"],
        ["T10", "案件查詢與 Flex/Carousel", "T08", "C"],
        ["T11", "WebGIS 示範頁", "T07,T08", "C"],
        ["T12", "Vendor API（pull）", "T08", "D"],
        ["T13", "PDF parser pipeline", "T08", "D"],
        ["T14", "整合測試與壓力測試", "T09,T10,T11,T12", "D"],
        ["T15", "上線部署與教育訓練", "T14", "D"],
    ]
    add_table_grid(doc, ["任務", "內容", "相依任務", "平行群組"], tasks)

    add_heading(doc, "12.2 平行化群組說明", 2)
    group_desc = [
        "群組 A：基礎啟動（需求、註冊、LINE 串接、狀態機骨架）。",
        "群組 B：資料核心（照片、標註、LRS、案件儲存）。",
        "群組 C：決策介面（審核、查詢、地圖）。",
        "群組 D：整合交付（外部 API、文件解析、測試、部署）。",
    ]
    for item in group_desc:
        add_paragraph_with_font(doc, f"- {item}")

    add_heading(doc, "12.3 技術堆疊摘要", 2)
    add_paragraph_with_font(
        doc,
        "後端以 FastAPI 為核心，前端地圖採 Leaflet.js，資料儲存採 JSON 檔案結構，影像與文件處理依賴 Pillow、PyMuPDF、pdfplumber、pytesseract，可在不導入資料庫前提下快速上線。",
    )

    add_heading(doc, "12.4 部署考量", 2)
    deployment = [
        "建議以 container 方式部署，區分 API、背景工作與靜態地圖服務。",
        "設定每日備份 case 目錄與 audit 檔案，採異地備援。",
        "API Key 與 LINE secrets 需放入安全設定，不寫入程式碼庫。",
        "建立監控：API 延遲、錯誤率、Webhook 重送率、儲存容量。",
    ]
    for item in deployment:
        add_paragraph_with_font(doc, f"- {item}")


def chapter_14(doc: Document) -> None:
    add_heading(doc, "第14章 定版功能與自動報告勾稽原則", 1)

    add_heading(doc, "14.1 自動產生報告勾稽原則（Word）", 2)
    section_14_1 = [
        "勾稽來源以照片標註（Evidence annotations.tags）為主。",
        "Table 1 採「先勾破壞模式，再勾致災原因」gate 原則，避免跨群誤勾。",
        "道路邊坡類：可勾道路上方邊坡滑動、道路下方邊坡滑動、整體性破壞，並聯動土質鬆軟、坡度過大。",
        "護岸/擋土牆類：已建立護岸、河道、上/下方擋土牆破壞對應規則，並聯動水文/排水/介面等致災原因。",
        "橋梁類：已建立橋墩、橋面、橋台、主梁等破壞模式與洪水沖刷/撞擊原因勾稽。",
        "Table 2 目前自動勾稽 Row1~Row5；Row6（熱危害）與 Row7（低溫危害）依現行決議維持手動不自動勾選。",
        "保留人工判斷欄位：Table 1 之「其他(請敘述)」與「致災原因需另辦理整體安全評估」。",
        "正式規格檔：docs/word_checkbox_mapping_spec.md，後續功能報告須納入本檔內容。",
    ]
    for item in section_14_1:
        add_paragraph_with_font(doc, f"- {item}")

    add_heading(doc, "14.2 LINE 與審核流程功能（定版）", 2)
    section_14_2 = [
        "全域選單 postback 於任何 flow 下均可正確路由，不再誤跳「請選擇審核類別」。",
        "查詢案件 -> 選狀態 -> 查看，已可正常開啟案件詳情。",
        "管理者於案件詳情可執行「通過 / 退回 / 結案」；一般使用者維持權限控管。",
        "標註互動優化：排除選項視為單選，點擊即自動前進到下一標註分類。",
        "多選標註採累積模式，完成所有選項後再按一次「確認選擇」進下一步，減少重複點擊。",
        "災害原因選單已加入「地震」並防止重複顯示。",
        "必要照片順序已統一為 P1 -> P2 -> P3 -> P4，並同步調整名稱與提示文案。",
    ]
    for item in section_14_2:
        add_paragraph_with_font(doc, f"- {item}")

    add_heading(doc, "14.3 WebGIS 與統計功能（定版）", 2)
    section_14_3 = [
        "WebGIS 圖面可視狀態統一包含：待審核、處理中、已結案、已退回。",
        "已退回案件若缺座標，使用工務段中心點作為低信心 fallback marker，以確保圖面可視。",
        "資訊摘要「最後更新」已改為完整日期時間（年月日+時分秒）。",
        "邊坡災害統計儀表板總數與各工務段件數已改為與 WebGIS 後台一致口徑。",
        "統一計數來源：僅納入可視案件狀態（pending_review / in_progress / closed / returned），排除 draft。",
        "WebGIS 後台刪除案件後，LINE 統計摘要與 stats 儀表板會同步更新，不再出現數量不一致。",
    ]
    for item in section_14_3:
        add_paragraph_with_font(doc, f"- {item}")

    add_heading(doc, "14.4 文件與穩定性說明", 2)
    add_paragraph_with_font(
        doc,
        "本次文件使用 python-docx 直接讀寫原始 .docx 並原地覆寫，避免非 Office 相容格式造成開啟錯誤。",
    )
    add_paragraph_with_font(
        doc,
        "若使用者端仍看到舊版內容，請先關閉檔案並重新開啟或清除暫存後再讀取最新文件。",
    )


def appendix_a(doc: Document) -> None:
    add_heading(doc, "附錄A 破壞模式與致災原因對照表", 1)
    add_heading(doc, "A.1 護岸/擋土牆類（9 組）", 2)
    retaining = [
        ["牆體龜裂", "長期滲水與排水失效"],
        ["牆面剝落", "材料劣化與凍融循環"],
        ["擋土牆傾斜", "基礎掏空或側向土壓增加"],
        ["牆趾淘刷", "豪雨沖刷與河道流速提升"],
        ["伸縮縫破壞", "溫差變形與維護不足"],
        ["排水孔阻塞", "淤積、落葉與漂流物堵塞"],
        ["護岸塊體位移", "基床流失與地盤鬆動"],
        ["護岸沉陷", "地基承載力下降"],
        ["護岸崩落", "連續降雨引發坡腳失穩"],
    ]
    add_table_grid(doc, ["破壞模式", "致災原因"], retaining)

    add_heading(doc, "A.2 道路邊坡類（10 組）", 2)
    slope = [
        ["表層滑動", "短延時強降雨"],
        ["深層滑動", "地層弱面活化"],
        ["落石", "岩體節理鬆動"],
        ["崩塌", "坡體飽和與剪力降低"],
        ["土石流入侵", "上游集水區崩塌"],
        ["路肩掏空", "坡腳沖刷"],
        ["路基下陷", "地下水位上升"],
        ["裂縫擴大", "地盤持續變位"],
        ["擋土設施失效", "排水不良與結構老化"],
        ["邊坡排水破壞", "排水溝堵塞溢流"],
    ]
    add_table_grid(doc, ["破壞模式", "致災原因"], slope)

    add_heading(doc, "A.3 橋梁類（9 組）", 2)
    bridge = [
        ["橋台沖刷", "河道洪峰流速過高"],
        ["橋墩裸露", "長期淘刷累積"],
        ["伸縮縫損壞", "溫差與重車荷載"],
        ["橋面破損", "材料疲勞與滲水"],
        ["護欄變形", "車撞或土石撞擊"],
        ["翼牆裂損", "地盤沉陷與側向壓力"],
        ["引道沉陷", "路基排水失效"],
        ["排水孔堵塞", "泥砂與落葉堆積"],
        ["下部結構裂縫", "反覆荷載與老化"],
    ]
    add_table_grid(doc, ["破壞模式", "致災原因"], bridge)


def appendix_b(doc: Document) -> None:
    add_heading(doc, "附錄B 工址環境調查項目", 1)
    add_paragraph_with_font(doc, "以下項目依輸入表單 Table 3 規劃，作為 Step 10 工址環境調查勾選清單。")
    add_table_grid(
        doc,
        ["分類", "調查項目", "說明"],
        [
            ["上邊坡", "落石", "觀察是否有新生落石、堆積或持續掉落跡象"],
            ["上邊坡", "崩塌", "評估崩塌範圍、厚度與是否持續擴大"],
            ["下邊坡", "路基缺口", "檢查路肩、坡腳是否出現缺口與掏空"],
            ["下邊坡", "下陷", "檢查路面沉陷、龜裂及車道通行安全"],
            ["結構物", "鋼筋裸露", "混凝土剝落導致鋼筋外露情形"],
            ["結構物", "危木", "邊坡危木、傾倒樹木對道路影響"],
            ["結構物", "護欄", "護欄變形、斷裂、基礎鬆動"],
            ["結構物", "坑洞", "路面坑洞大小、深度與分布"],
            ["結構物", "電桿", "電桿傾斜、基礎鬆動與纜線風險"],
            ["橋梁河道", "鄰河", "河道距離、沖刷影響與側蝕風險"],
            ["橋梁河道", "曲流", "彎道主流衝擊點與橋梁安全關聯"],
            ["其他", "天氣", "當下天候與過去 24 小時降雨描述"],
            ["其他", "降雪", "高山路段降雪、融雪與路面結冰風險"],
        ],
    )


def appendix_c(doc: Document) -> None:
    add_heading(doc, "附錄C Case JSON Schema", 1)
    schema_text = """{
  "case_id": "case_20260228_0001",
  "district": {
    "district_id": "jingmei",
    "district_name": "景美工務段"
  },
  "reporter": {
    "user_id": "Uxxxxxxxx",
    "display_name": "巡查員A",
    "real_name": "王大明",
    "role": "user"
  },
  "route": {
    "route_id": "TH9",
    "route_name": "台9",
    "direction": "北上",
    "road_section": "台9 45k+200"
  },
  "location": {
    "lat": 24.987654,
    "lng": 121.543210,
    "source": "lrs",
    "milepost": "45k+200",
    "lrs_confidence": 0.91,
    "candidate_mileposts": ["45k+180", "45k+200", "45k+220"]
  },
  "damage": {
    "mode_category": "道路邊坡",
    "mode": "崩塌",
    "cause": "坡體飽和與剪力降低",
    "description": "外側車道受土石堆積影響，需立即清除。"
  },
  "photos": [
    {
      "photo_id": "sha256:...",
      "photo_type": "P1",
      "required": true,
      "captured_at": "2026-02-28T09:10:00+08:00",
      "annotation": {
        "severity": "S2",
        "tags": ["北向", "陰天", "能見度中"],
        "custom_notes": ["雨後路面濕滑"],
        "ai_ready": {
          "verified": true,
          "quality_score": 0.88
        }
      }
    }
  ],
  "survey": {
    "upper_slope": ["崩塌"],
    "lower_slope": ["路基缺口"],
    "structures": ["護欄"],
    "bridge_channel": [],
    "others": ["天氣"]
  },
  "cost_estimate": {
    "enabled": false,
    "amount_range": null,
    "currency": "TWD"
  },
  "status": {
    "technical_track": "milepost_resolved",
    "business_track": "pending_review"
  },
  "timeline": {
    "created_at": "2026-02-28T09:08:00+08:00",
    "submitted_at": "2026-02-28T09:25:00+08:00",
    "updated_at": "2026-02-28T09:26:30+08:00",
    "closed_at": null
  },
  "audit_ref": "./audit.jsonl",
  "attachments": {
    "evidence_dir": "./evidence",
    "derived_dir": "./derived",
    "thumbnails_dir": "./thumbnails"
  }
}"""
    add_paragraph_with_font(doc, schema_text)
    add_paragraph_with_font(
        doc,
        "Schema 實作重點：欄位採明確型別、允許 null 的欄位需有業務理由、所有時間欄位使用 ISO 8601、列舉值維持固定集合避免髒資料。",
    )


def appendix_d(doc: Document) -> None:
    add_heading(doc, "附錄D API 端點列表", 1)
    endpoints = [
        ["POST", "/line/webhook", "接收 LINE 事件並觸發流程"],
        ["GET", "/health", "服務健康檢查"],
        ["POST", "/auth/register", "使用者註冊資料提交"],
        ["POST", "/auth/approve/{user_id}", "決策人員核准帳號"],
        ["POST", "/auth/reject/{user_id}", "決策人員退回註冊"],
        ["GET", "/users/me", "查詢目前使用者資料"],
        ["POST", "/cases", "建立案件草稿"],
        ["PUT", "/cases/{case_id}/step/{step_no}", "更新指定步驟資料"],
        ["POST", "/cases/{case_id}/photos", "上傳照片與 metadata"],
        ["PUT", "/cases/{case_id}/photos/{photo_id}/annotation", "更新照片標註"],
        ["POST", "/cases/{case_id}/submit", "確認送出案件"],
        ["GET", "/cases/my", "使用者查詢個人案件"],
        ["GET", "/cases", "決策人員條件查詢案件"],
        ["GET", "/cases/{case_id}", "查詢單一案件完整內容"],
        ["PUT", "/cases/{case_id}/business-status", "更新業務狀態"],
        ["POST", "/cases/{case_id}/return", "退回補件"],
        ["POST", "/cases/{case_id}/close", "結案"],
        ["GET", "/gis/cases", "WebGIS 案件點位資料"],
        ["GET", "/vendor/cases", "Vendor API 增量拉取案件"],
        ["GET", "/vendor/cases/{case_id}", "Vendor API 單案查詢"],
        ["DELETE", "/api/v1/cases/{case_id}", "刪除案件（含稽核紀錄與 LINE 通知）"],
        ["GET", "/api/v1/cases/{case_id}/word", "下載案件 Word 報表（公路災害工程內容概述表）"],
        ["GET", "/api/v1/cases/{case_id}/photos/{photo_id}/thumbnail", "取得照片縮圖"],
        ["POST", "/lrs/resolve", "座標轉里程與信心分數"],
        ["POST", "/pdf/parse", "PDF 解析任務"],
        ["GET", "/stats/overview", "決策儀表板統計資料"],
    ]
    add_table_grid(doc, ["Method", "Path", "Description"], endpoints)

    add_paragraph_with_font(
        doc,
        "所有 API 回應均採 UTF-8 JSON、統一錯誤格式（error_code、message、trace_id），並以 OpenAPI 文件提供測試與整合參考。",
    )


def build_document() -> None:
    doc = Document()
    configure_styles(doc)

    add_cover_page(doc)
    add_page_break(doc)
    add_toc_placeholder(doc)

    chapters = [
        chapter_1,
        chapter_2,
        chapter_3,
        chapter_4,
        chapter_5,
        chapter_6,
        chapter_7,
        chapter_8,
        chapter_9,
        chapter_10,
        chapter_11,
        chapter_12,
        chapter_14,
    ]

    for idx, chapter in enumerate(chapters):
        add_page_break(doc)
        chapter(doc)
        if idx < len(chapters) - 1:
            add_paragraph_with_font(doc, "")

    add_page_break(doc)
    appendix_a(doc)
    add_page_break(doc)
    appendix_b(doc)
    add_page_break(doc)
    appendix_c(doc)
    add_page_break(doc)
    appendix_d(doc)

    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT_PATH))
    _postprocess_docx(OUTPUT_PATH)


def _postprocess_docx(path: Path) -> None:
    """Clean up the docx ZIP for maximum Word compatibility.

    python-docx's default template ships with stylesWithEffects.xml
    (a Word 2010 artifact) and sets compatibilityMode=14.  Newer Word
    versions may refuse to open such files or flag them as corrupted.

    This function:
    1. Removes word/stylesWithEffects.xml
    2. Strips the corresponding Content_Types override and relationship
    3. Bumps compatibilityMode from 14 to 15 (Word 2013+)
    """
    import xml.etree.ElementTree as ET
    import zipfile as zf

    tmp = path.with_suffix(".docx.tmp")
    with zf.ZipFile(str(path), "r") as zin, zf.ZipFile(str(tmp), "w", zf.ZIP_DEFLATED) as zout:
        for name in zin.namelist():
            data = zin.read(name)

            # 1. Drop stylesWithEffects.xml entirely
            if "stylesWithEffects" in name:
                continue

            # 2. Remove its Content_Types entry
            if name == "[Content_Types].xml":
                data = data.replace(
                    b'<Override PartName="/word/stylesWithEffects.xml"'
                    b' ContentType="application/vnd.ms-word.stylesWithEffects+xml"/>',
                    b"",
                )

            # 3. Remove its relationship entry
            if name == "word/_rels/document.xml.rels":
                ns = "http://schemas.openxmlformats.org/package/2006/relationships"
                root = ET.fromstring(data)
                for rel in list(root.findall(f"{{{ns}}}Relationship")):
                    if "stylesWithEffects" in (rel.get("Target") or ""):
                        root.remove(rel)
                data = ET.tostring(root, xml_declaration=True, encoding="UTF-8")

            # 4. Bump compatibilityMode to 15
            if name == "word/settings.xml":
                data = data.replace(
                    b'w:name="compatibilityMode"'
                    b' w:uri="http://schemas.microsoft.com/office/word"'
                    b' w:val="14"',
                    b'w:name="compatibilityMode"'
                    b' w:uri="http://schemas.microsoft.com/office/word"'
                    b' w:val="15"',
                )

            zout.writestr(name, data)

    # Atomic replace
    tmp.replace(path)


def verify_document() -> None:
    doc = Document(str(OUTPUT_PATH))
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

    required_phrases = [
        "邊坡災害通報與資訊整合管理系統",
        "版本：v1.1.0",
        "第1章 專案概述",
        "第6章 災害通報流程（12步驟）",
        "第7章 照片類型與標註系統",
        "7.1 共用照片類型（common）",
        "Step 3：省道里程輸入",
        "10.5 Word 報表自動產生",
        "第14章 定版功能與自動報告勾稽原則",
        "附錄D API 端點列表",
        "交通部公路局北區養護工程分局",
    ]
    for phrase in required_phrases:
        if not any(phrase in p for p in paragraphs):
            raise ValueError(f"驗證失敗：找不到文字：{phrase}")

    mojibake_markers = ["Ã", "å", "ç", "è", "æ", "�"]
    suspicious = [p for p in paragraphs if any(marker in p for marker in mojibake_markers)]
    if suspicious:
        raise ValueError("驗證失敗：疑似編碼錯誤字元出現在段落中。")

    spot_checks = paragraphs[:25]
    if len(spot_checks) < 10:
        raise ValueError("驗證失敗：段落數不足，無法完成 10 段以上抽查。")

    print("文件已產生：", OUTPUT_PATH)
    print("抽查段落（前 12 段）：")
    for idx, text in enumerate(spot_checks[:12], start=1):
        print(f"{idx:02d}. {text}")
    print(f"總段落數：{len(paragraphs)}")


if __name__ == "__main__":
    build_document()
    verify_document()
