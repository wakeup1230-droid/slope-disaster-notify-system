#!/usr/bin/env python3
# pyright: reportUnusedImport=false, reportUnknownVariableType=false, reportUnknownParameterType=false, reportMissingParameterType=false, reportUnknownMemberType=false, reportUnknownArgumentType=false, reportGeneralTypeIssues=false, reportMissingTypeArgument=false, reportImplicitStringConcatenation=false
"""
更新規劃文件至 v1.1.0：
1) 更新版本號與日期
2) 以 photo_tags.json 動態重建第7章「照片類型與標註系統」
3) 覆寫原始 docx（不另存備份）
"""

from __future__ import annotations

import json
import re
from datetime import datetime
from pathlib import Path
from typing import cast

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

BASE_DIR = Path(__file__).resolve().parent
DOCX_PATH = BASE_DIR / "邊坡災害通報與資訊整合管理系統_規劃文件.docx"
PHOTO_TAGS_PATH = BASE_DIR / "app" / "data" / "photo_tags.json"

TARGET_VERSION = "v1.1.0"

FONT_NAME = "微軟正黑體"
FONT_SIZE_NORMAL = Pt(12)
FONT_SIZE_H1 = Pt(15)
FONT_SIZE_H2 = Pt(14)
FONT_SIZE_H3 = Pt(13)

CATEGORY_LAYOUT = [
    ("common", "7.1 共用照片類型（common）"),
    ("optional", "7.2 選用照片類型（optional）"),
    ("revetment_retaining", "7.3 護岸/擋土牆專用照片類型（revetment_retaining）"),
    ("road_slope", "7.4 道路邊坡專用照片類型（road_slope）"),
    ("bridge", "7.5 橋梁專用照片類型（bridge）"),
]


# helpers (reused pattern)

def set_run(run, text, font_name=FONT_NAME, font_size=FONT_SIZE_NORMAL, bold=False):
    """Apply standard formatting to a run."""
    run.text = text
    run.font.name = font_name
    run.font.size = font_size
    run.font.bold = bold
    rpr = run._element.get_or_add_rPr()
    rFonts = rpr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), font_name)


def make_paragraph(doc, text, style="Normal", font_size=FONT_SIZE_NORMAL, bold=False):
    """Create a new paragraph with standard formatting."""
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    set_run(run, text, font_size=font_size, bold=bold)
    return p


def find_paragraph_index(doc, text_contains):
    """Find the index of the first paragraph containing given text."""
    for i, p in enumerate(doc.paragraphs):
        if text_contains in p.text:
            return i
    return -1


def insert_paragraph_after(doc, ref_para, text, style="Normal", font_size=FONT_SIZE_NORMAL, bold=False):
    """Insert a new paragraph after a reference paragraph."""
    new_p = doc.add_paragraph(style=style)
    run = new_p.add_run(text)
    set_run(run, text, font_size=font_size, bold=bold)
    ref_para._element.addnext(new_p._element)
    return new_p


def add_table_grid(table):
    """Apply Table Grid borders to an existing table."""
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = qn(f"w:{edge}")
        border = borders.find(tag)
        if border is None:
            border = OxmlElement(f"w:{edge}")
            borders.append(border)
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "8")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "000000")


def _replace_paragraph_text(para, new_text, font_size=FONT_SIZE_NORMAL, bold=False):
    for run in para.runs:
        run.text = ""
    if para.runs:
        set_run(para.runs[0], new_text, font_size=font_size, bold=bold)
    else:
        run = para.add_run(new_text)
        set_run(run, new_text, font_size=font_size, bold=bold)


def _is_chapter_heading(text: str, number: int) -> bool:
    cleaned = text.strip().replace(" ", "")
    if not cleaned:
        return False
    patterns = [
        rf"^第{number}章",
        rf"^{number}章",
        rf"^{number}[\.．、)]",
        rf"^Chapter{number}",
    ]
    return any(re.search(pattern, cleaned, flags=re.IGNORECASE) for pattern in patterns)


def _find_chapter_index(doc, chapter_number: int) -> int:
    for idx, para in enumerate(doc.paragraphs):
        if _is_chapter_heading(para.text, chapter_number):
            return idx
    return -1


def _insert_lines_after(doc, anchor_para, lines):
    last = anchor_para
    inserted = 0
    for line in lines:
        text = line.get("text", "")
        style = line.get("style", "Normal")
        font_size = line.get("font_size", FONT_SIZE_NORMAL)
        bold = line.get("bold", False)
        align = line.get("align")
        last = insert_paragraph_after(doc, last, text, style=style, font_size=font_size, bold=bold)
        if align is not None:
            last.alignment = align
        inserted += 1
    return last, inserted


def _delete_paragraph_range(doc, start_idx: int, end_idx: int) -> int:
    if start_idx < 0 or end_idx < 0 or end_idx <= start_idx:
        return 0
    body = doc.element.body
    targets = [doc.paragraphs[i]._element for i in range(start_idx, end_idx)]
    for elem in targets:
        body.remove(elem)
    return len(targets)


def _tag_line(tag_def: dict) -> str:
    mode = "多選" if tag_def.get("multi_select") else "單選"
    tier = tag_def.get("tier", "unknown")
    return f"- {tag_def.get('category_name', '未命名分類')}（{mode}，tier: {tier}）"


def _options_line(prefix: str, tags) -> str:
    if not tags:
        return f"  {prefix}：無"
    labels = "、".join(tag.get("label", "") for tag in tags)
    return f"  {prefix}：{labels}"


def _build_chapter7_lines(photo_tags_data):
    lines: list[dict] = []
    stats = {
        "category_count": 0,
        "photo_type_count": 0,
        "photo_tag_category_count": 0,
        "judgment_tag_category_count": 0,
        "option_count": 0,
        "exclusion_count": 0,
    }

    lines.append(
        {
            "text": "第7章 照片類型與標註系統",
            "style": "Heading 1",
            "font_size": FONT_SIZE_H1,
            "bold": True,
        }
    )
    lines.append(
        {
            "text": "本系統共定義 10 種照片類型（P1-P10），分布於 5 大分類（common、optional、revetment_retaining、road_slope、bridge），以下完整列出各分類之照片類型、標註分類與全部可選項目。",
            "style": "Normal",
            "font_size": FONT_SIZE_NORMAL,
            "bold": False,
        }
    )

    for category_key, category_heading in CATEGORY_LAYOUT:
        category_block = photo_tags_data.get(category_key, {})
        stats["category_count"] += 1
        lines.append(
            {
                "text": category_heading,
                "style": "Heading 2",
                "font_size": FONT_SIZE_H2,
                "bold": True,
            }
        )

        for photo_type_key, photo_type_def in category_block.items():
            required = "必填" if photo_type_def.get("required") else "選填"
            max_photos = photo_type_def.get("max_photos", "-")
            heading_text = (
                f"{photo_type_key} {photo_type_def.get('name', '')}"
                f"（{required}，最多 {max_photos} 張）"
            )
            lines.append(
                {
                    "text": heading_text,
                    "style": "Heading 3",
                    "font_size": FONT_SIZE_H3,
                    "bold": True,
                }
            )
            stats["photo_type_count"] += 1

            lines.append(
                {
                    "text": "照片標註項目 (photo_tags)",
                    "style": "Normal",
                    "font_size": FONT_SIZE_NORMAL,
                    "bold": True,
                }
            )
            photo_tag_groups = photo_type_def.get("photo_tags", [])
            if not photo_tag_groups:
                lines.append({"text": "- 無", "style": "Normal"})
            for tag_group in photo_tag_groups:
                lines.append({"text": _tag_line(tag_group), "style": "Normal"})
                lines.append({"text": _options_line("選項", tag_group.get("tags", [])), "style": "Normal"})
                if tag_group.get("exclusion_tags"):
                    lines.append(
                        {
                            "text": _options_line("排除選項", tag_group.get("exclusion_tags", [])),
                            "style": "Normal",
                        }
                    )
                stats["photo_tag_category_count"] += 1
                stats["option_count"] += len(tag_group.get("tags", []))
                stats["exclusion_count"] += len(tag_group.get("exclusion_tags", []))

            lines.append(
                {
                    "text": "人工研判項目 (judgment_tags)",
                    "style": "Normal",
                    "font_size": FONT_SIZE_NORMAL,
                    "bold": True,
                }
            )
            judgment_tag_groups = photo_type_def.get("judgment_tags", [])
            if not judgment_tag_groups:
                lines.append({"text": "- 無", "style": "Normal"})
            for tag_group in judgment_tag_groups:
                base_line = _tag_line(tag_group)
                if tag_group.get("input_type") == "text":
                    base_line += "（自由文字輸入）"
                lines.append({"text": base_line, "style": "Normal"})
                if tag_group.get("input_type") == "text":
                    lines.append({"text": "  選項：自由文字輸入", "style": "Normal"})
                else:
                    lines.append({"text": _options_line("選項", tag_group.get("tags", [])), "style": "Normal"})
                if tag_group.get("exclusion_tags"):
                    lines.append(
                        {
                            "text": _options_line("排除選項", tag_group.get("exclusion_tags", [])),
                            "style": "Normal",
                        }
                    )
                stats["judgment_tag_category_count"] += 1
                stats["option_count"] += len(tag_group.get("tags", []))
                stats["exclusion_count"] += len(tag_group.get("exclusion_tags", []))

    return lines, stats


def update_version(doc) -> int:
    updated = 0
    version_pattern = re.compile(r"v\d+\.\d+\.\d+", flags=re.IGNORECASE)
    for para in doc.paragraphs:
        text = para.text.strip()
        if "版本" not in text and "Version" not in text:
            continue
        if version_pattern.search(text):
            new_text = version_pattern.sub(TARGET_VERSION, text)
        elif "版本" in text:
            new_text = "版本：v1.1.0"
        else:
            new_text = "Version: v1.1.0"
        if new_text != text:
            _replace_paragraph_text(para, new_text)
            updated += 1
    return updated


def update_date(doc) -> int:
    today_zh = datetime.now().strftime("%Y年%m月%d日")
    updated = 0
    date_pattern_zh = re.compile(r"\d{4}年\d{1,2}月\d{1,2}日")
    date_pattern_iso = re.compile(r"\d{4}[-/]\d{1,2}[-/]\d{1,2}")

    for para in doc.paragraphs:
        text = para.text.strip()
        if "日期" not in text and "Date" not in text:
            continue

        if date_pattern_zh.search(text):
            new_text = date_pattern_zh.sub(today_zh, text)
        elif date_pattern_iso.search(text):
            new_text = date_pattern_iso.sub(today_zh, text)
        elif "日期" in text:
            new_text = f"日期：{today_zh}"
        else:
            new_text = f"Date: {today_zh}"

        if new_text != text:
            _replace_paragraph_text(para, new_text)
            updated += 1

    return updated


def replace_chapter7(doc, photo_tags_data) -> dict:
    chapter7_idx = _find_chapter_index(doc, 7)
    chapter8_idx = _find_chapter_index(doc, 8)

    lines, data_stats = _build_chapter7_lines(photo_tags_data)

    deleted_count = 0
    insert_before_ch8 = False
    appended_at_end = False

    if chapter7_idx >= 0:
        if chapter8_idx > chapter7_idx:
            deleted_count = _delete_paragraph_range(doc, chapter7_idx, chapter8_idx)
        else:
            deleted_count = _delete_paragraph_range(doc, chapter7_idx, len(doc.paragraphs))

    elif chapter8_idx >= 0:
        insert_before_ch8 = True

    else:
        appended_at_end = True

    if insert_before_ch8:
        chapter8_idx = _find_chapter_index(doc, 8)
        if chapter8_idx <= 0:
            anchor = doc.paragraphs[0]
            first_line = lines[0]
            new_p = doc.add_paragraph(style=first_line.get("style", "Normal"))
            run = new_p.add_run(first_line.get("text", ""))
            set_run(
                run,
                first_line.get("text", ""),
                font_size=first_line.get("font_size", FONT_SIZE_NORMAL),
                bold=first_line.get("bold", False),
            )
            anchor._element.addprevious(new_p._element)
            _, inserted_count = _insert_lines_after(doc, new_p, lines[1:])
            inserted_count += 1
        else:
            anchor = doc.paragraphs[chapter8_idx - 1]
            _, inserted_count = _insert_lines_after(doc, anchor, lines)
    elif appended_at_end:
        anchor = doc.paragraphs[-1]
        _, inserted_count = _insert_lines_after(doc, anchor, lines)
    else:
        chapter7_idx = _find_chapter_index(doc, 7)
        if chapter7_idx > 0:
            anchor = doc.paragraphs[chapter7_idx - 1]
            _, inserted_count = _insert_lines_after(doc, anchor, lines)
        else:
            anchor = doc.paragraphs[0]
            first_line = lines[0]
            new_p = doc.add_paragraph(style=first_line.get("style", "Normal"))
            run = new_p.add_run(first_line.get("text", ""))
            set_run(
                run,
                first_line.get("text", ""),
                font_size=first_line.get("font_size", FONT_SIZE_NORMAL),
                bold=first_line.get("bold", False),
            )
            anchor._element.addprevious(new_p._element)
            _, inserted_count = _insert_lines_after(doc, new_p, lines[1:])
            inserted_count += 1

    return {
        "chapter7_found": chapter7_idx >= 0,
        "chapter8_found": chapter8_idx >= 0,
        "deleted_paragraphs": deleted_count,
        "inserted_paragraphs": inserted_count,
        **data_stats,
    }


def main():
    print(f"[INFO] Loading docx: {DOCX_PATH}")
    doc = Document(str(DOCX_PATH))

    print(f"[INFO] Loading tags: {PHOTO_TAGS_PATH}")
    with PHOTO_TAGS_PATH.open("r", encoding="utf-8") as f:
        photo_tags_data = cast(dict, json.load(f))

    version_updates = update_version(doc)
    date_updates = update_date(doc)
    chapter_stats = replace_chapter7(doc, photo_tags_data)

    print(f"[INFO] Saving updated docx in-place: {DOCX_PATH}")
    doc.save(str(DOCX_PATH))

    print("[DONE] 更新完成")
    print(f"- 版本欄位更新數：{version_updates}")
    print(f"- 日期欄位更新數：{date_updates}")
    print(
        "- 第7章重建："
        f"刪除 {chapter_stats['deleted_paragraphs']} 段、"
        f"新增 {chapter_stats['inserted_paragraphs']} 段"
    )
    print(
        "- 標註資料統計："
        f"分類 {chapter_stats['category_count']}、"
        f"照片類型 {chapter_stats['photo_type_count']}、"
        f"photo_tags 分類 {chapter_stats['photo_tag_category_count']}、"
        f"judgment_tags 分類 {chapter_stats['judgment_tag_category_count']}、"
        f"選項 {chapter_stats['option_count']}、"
        f"排除選項 {chapter_stats['exclusion_count']}"
    )


if __name__ == "__main__":
    main()
