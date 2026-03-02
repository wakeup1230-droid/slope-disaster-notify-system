"""Word document generator service."""

from __future__ import annotations

import io
import re
from pathlib import Path
from typing import Any

from docx import Document as load_document
from docx.document import Document as DocxDocument
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from docx.text.paragraph import Paragraph
from docx.text.run import Run

from app.core.logging_config import get_logger
from app.models.case import Case, CostBreakdownItem
from app.models.evidence import EvidenceManifest

logger = get_logger(__name__)

DEFAULT_TEMPLATE = Path("Input/公路災害工程內容概述表_空白.docx")


class WordGenerator:
    """Generate a filled Word document from case data."""

    COMPLETENESS_FIELDS: list[tuple[str, str, bool]] = [
        # (field_key, display_name, is_required)
        ("reporting_year", "年度", True),
        ("disaster_type", "災害類型", True),
        ("processing_type", "處理類型", True),
        ("project_name", "工程名稱", True),
        ("disaster_date", "災害日期", True),
        ("town_village", "地點(鄉鎮里)", True),
        ("nearby_landmark", "鄰近地標", False),
        ("repeat_disaster", "重複致災", True),
        ("coordinates", "座標", True),
        ("damage_mode", "破壞模式", True),
        ("damage_cause", "致災原因", True),
        ("description", "災損描述", True),
        ("location_map", "位置圖", True),
        ("photos", "照片", True),
        ("original_protection", "原設計保護型式", False),
        ("analysis_review", "分析與檢討", False),
        ("estimated_cost", "初估經費", True),
        ("cost_breakdown", "經費明細", True),
        ("design_docs", "設計圖說", False),
        ("soil_conservation", "水土保持", True),
        ("safety_assessment", "安全評估", False),
        ("site_survey", "工址環境危害", True),
        ("other_supplement", "其他補充", False),
        ("reporting_agency", "提報機關", True),
        ("created_by", "填報人", True),
    ]

    @staticmethod
    def calculate_completeness(case: Case) -> dict[str, Any]:
        """計算Word文檔位字完整度。
        
        Returns: {
            "filled": int,
            "total": int,
            "percentage": int,
            "missing": [{"key": str, "name": str, "required": bool}],
        }
        """
        def _is_filled(key: str) -> bool:
            match key:
                case "town_village":
                    return bool(case.town_name or case.village_name)
                case "coordinates" | "location_map":
                    return case.primary_coordinate is not None
                case "damage_mode":
                    return bool(case.damage_mode_name)
                case "damage_cause":
                    return bool(case.damage_cause_names)
                case "photos":
                    return case.photo_count > 0
                case "cost_breakdown":
                    return bool(case.cost_breakdown)
                case "design_docs":
                    return bool(case.design_doc_evidence_id)
                case "site_survey":
                    return bool(case.site_survey)
                case "created_by":
                    return bool(case.created_by and case.created_by.real_name)
                case _:
                    return bool(getattr(case, key, None))

        filled = 0
        missing = []
        for key, name, required in WordGenerator.COMPLETENESS_FIELDS:
            if _is_filled(key):
                filled += 1
            else:
                missing.append({"key": key, "name": name, "required": required})
        total = len(WordGenerator.COMPLETENESS_FIELDS)
        return {
            "filled": filled,
            "total": total,
            "percentage": round(filled * 100 / total) if total > 0 else 0,
            "missing": missing,
        }


    def __init__(self, template_path: Path = DEFAULT_TEMPLATE, cases_dir: Path | None = None) -> None:
        self._template_path: Path = template_path
        self._cases_dir: Path = cases_dir or Path("storage/cases")
        if not self._template_path.exists():
            raise FileNotFoundError(f"Template not found: {self._template_path}")

    def generate(self, case: Case, manifest: EvidenceManifest | None = None) -> bytes:
        """Generate and return the filled document as bytes."""
        doc = load_document(str(self._template_path))

        self._fill_header(doc, case)
        self._fill_basic_info(doc, case)
        self._fill_coordinates(doc, case)
        self._fill_description(doc, case)
        self._fill_location_map(doc, case)
        self._fill_photos(doc, case, manifest)
        self._fill_analysis(doc, case, manifest)
        self._fill_cost(doc, case)
        self._fill_other_supplement(doc, case)
        self._fill_soil_conservation(doc, case)
        self._fill_national_park(doc, case)
        self._fill_hazard_table(doc, case, manifest)
        self._fill_signatures(doc, case)

        buffer = io.BytesIO()
        doc.save(buffer)
        return buffer.getvalue()

    @staticmethod
    def _set_run_font(run: Run, font_name: str = "標楷體", size_pt: int = 14) -> None:
        """Apply template font to a run with red color."""
        run.font.name = font_name
        run.font.size = Pt(size_pt)
        run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
        rpr = run._element.get_or_add_rPr()
        r_fonts = rpr.find(qn("w:rFonts"))
        if r_fonts is None:
            r_fonts = OxmlElement("w:rFonts")
            rpr.insert(0, r_fonts)
        r_fonts.set(qn("w:ascii"), font_name)
        r_fonts.set(qn("w:hAnsi"), font_name)
        r_fonts.set(qn("w:eastAsia"), font_name)

    @classmethod
    def _set_paragraph_text(cls, paragraph: Paragraph, text: str) -> None:
        """Replace paragraph text and apply font to inserted run."""
        paragraph.clear()
        run = paragraph.add_run(text)
        cls._set_run_font(run)

    @classmethod
    def _append_text_to_paragraph(cls, paragraph: Paragraph, text: str) -> None:
        """Append text to a paragraph using template font."""
        run = paragraph.add_run(text)
        cls._set_run_font(run)

    @classmethod
    def _insert_paragraph_after(cls, paragraph: Paragraph, text: str) -> Paragraph:
        """Insert a paragraph after the given paragraph."""
        new_p = OxmlElement("w:p")
        paragraph._element.addnext(new_p)
        new_paragraph = Paragraph(new_p, paragraph._parent)
        if text:
            run = new_paragraph.add_run(text)
            cls._set_run_font(run)
        return new_paragraph

    @staticmethod
    def _find_paragraph_containing(doc: DocxDocument, text: str) -> Paragraph | None:
        """Find first paragraph containing text."""
        for paragraph in doc.paragraphs:
            if text in paragraph.text:
                return paragraph
        return None

    @staticmethod
    def _find_paragraph_index_containing(doc: DocxDocument, text: str) -> int | None:
        """Find index of first paragraph containing text."""
        for idx, paragraph in enumerate(doc.paragraphs):
            if text in paragraph.text:
                return idx
        return None

    @staticmethod
    def _find_table_containing(doc: DocxDocument, text: str) -> Any | None:
        """Find first table containing target text in any cell."""
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if text in cell.text:
                        return table
        return None

    @staticmethod
    def _extract_photo_tags(manifest: EvidenceManifest | None) -> dict[str, dict[str, set[str]]]:
        """Extract aggregated photo tags from manifest, grouped by photo_type and category.

        Returns:
            {
                "P1": {"site_risks": {"upslope_rockfall", ...}, "weather": {...}},
                "P2": {"visible_damage": {"rockfall", ...}, "slope_location": {...}},
                ...
            }
        """
        result: dict[str, dict[str, set[str]]] = {}
        if not manifest or not manifest.evidence:
            return result
        for ev in manifest.evidence:
            ptype = ev.photo_type
            if not ptype:
                continue
            if not ev.annotations or not ev.annotations.tags:
                continue
            if ptype not in result:
                result[ptype] = {}
            for tag in ev.annotations.tags:
                cat = tag.category
                tid = tag.tag_id
                if cat not in result[ptype]:
                    result[ptype][cat] = set()
                result[ptype][cat].add(tid)
        return result
    @staticmethod
    def _parse_date(date_value: str) -> tuple[str, str, str]:
        """Parse date from YYYY-MM-DD or 民國年格式."""
        date_value = date_value.strip()
        if not date_value:
            return "", "", ""

        if "T" in date_value:
            date_value = date_value.split("T", 1)[0]
        date_value = date_value.replace("/", "-").replace(".", "-")

        if "-" in date_value:
            parts = date_value.split("-")
            if len(parts) == 3:
                try:
                    year = int(parts[0])
                    month = int(parts[1])
                    day = int(parts[2])
                    roc_year = year - 1911 if year >= 1912 else year
                    return str(roc_year), str(month), str(day)
                except ValueError:
                    return "", "", ""

        match = re.search(r"(\d+)\s*年\s*(\d+)\s*月\s*(\d+)\s*日", date_value)
        if match:
            return match.group(1), match.group(2), match.group(3)
        return "", "", ""

    @staticmethod
    def _format_number(value: float | None) -> str:
        """Format number for display in cost section."""
        if value is None:
            return ""
        if float(value).is_integer():
            return f"{int(value):,}"
        return f"{value:,.2f}"

    def _fill_header(self, doc: DocxDocument, case: Case) -> None:
        """Fill P1 year and P2 disaster/processing type."""
        if len(doc.paragraphs) <= 2:
            return

        p1 = doc.paragraphs[1]
        if case.reporting_year and p1.runs:
            p1.runs[0].text = f"  {case.reporting_year}  "
            self._set_run_font(p1.runs[0])

        p2 = doc.paragraphs[2]
        if len(p2.runs) < 6:
            return

        if case.disaster_type == "一般":
            p2.runs[0].text = "■一般災害  "
            p2.runs[1].text = "□專案災害"
        elif case.disaster_type == "專案":
            p2.runs[0].text = "□一般災害  "
            p2.runs[1].text = "■專案災害"

        if case.processing_type == "搶修":
            p2.runs[4].text = "■搶修 "
            p2.runs[5].text = "□復建 經費明細表"
        elif case.processing_type == "復建":
            p2.runs[4].text = "□搶修 "
            p2.runs[5].text = "■復建 經費明細表"

        for idx in (0, 1, 4, 5):
            self._set_run_font(p2.runs[idx])

    def _fill_basic_info(self, doc: DocxDocument, case: Case) -> None:
        """Fill P4-P8 basic case information."""
        if len(doc.paragraphs) <= 8:
            return

        p4 = doc.paragraphs[4]
        if case.project_name:
            p4.runs[0].text = f"工程名稱：{case.project_name}"
            self._set_run_font(p4.runs[0])

        p5 = doc.paragraphs[5]
        if case.disaster_date:
            year, month, day = self._parse_date(case.disaster_date)
            if len(p5.runs) >= 7 and year and month and day:
                p5.runs[1].text = year
                p5.runs[3].text = month
                p5.runs[5].text = day
                for idx in (1, 3, 5):
                    self._set_run_font(p5.runs[idx])

        p6 = doc.paragraphs[6]
        if len(p6.runs) >= 6:
            if case.town_name:
                p6.runs[1].text = case.town_name
                p6.runs[2].text = ""
                self._set_run_font(p6.runs[1])
            if case.village_name:
                p6.runs[4].text = case.village_name
                self._set_run_font(p6.runs[4])

        p7 = doc.paragraphs[7]
        if case.nearby_landmark and p7.runs:
            p7.runs[0].text = f"  {case.nearby_landmark}"
            self._set_run_font(p7.runs[0])
            for run in p7.runs[1:]:
                run.text = ""

        p8 = doc.paragraphs[8]
        if len(p8.runs) >= 6:
            if case.repeat_disaster == "否":
                p8.runs[1].text = "■否"
                p8.runs[3].text = "□是　"
            elif case.repeat_disaster == "是":
                p8.runs[1].text = "□否"
                p8.runs[3].text = "■是　"
                if case.repeat_disaster_year:
                    p8.runs[4].text = f" {case.repeat_disaster_year} "
                    self._set_run_font(p8.runs[4])
            for idx in (1, 3):
                self._set_run_font(p8.runs[idx])

    def _fill_coordinates(self, doc: DocxDocument, case: Case) -> None:
        """Coordinates are now appended after photos by _fill_photos. No-op."""
        pass

    def _fill_description(self, doc: DocxDocument, case: Case) -> None:
        """Fill P12 and P14 appended paragraphs."""
        if case.description and len(doc.paragraphs) > 12:
            self._insert_paragraph_after(doc.paragraphs[12], case.description)

        if case.original_protection:
            paragraph = self._find_paragraph_containing(doc, "原設計保護型式")
            if paragraph is not None:
                self._insert_paragraph_after(paragraph, case.original_protection)

    def _fill_location_map(self, doc: DocxDocument, case: Case) -> None:
        """Generate a static map image and insert after P16 (位置簡圖)."""
        if not case.primary_coordinate:
            return

        map_bytes = self._generate_static_map(
            case.primary_coordinate.lat,
            case.primary_coordinate.lon,
        )
        if not map_bytes:
            return

        paragraph = self._find_paragraph_containing(doc, "位置簡圖")
        if paragraph is None:
            return

        from docx.shared import Cm

        img_paragraph = self._insert_paragraph_after(paragraph, "")
        run = img_paragraph.add_run()
        run.add_picture(io.BytesIO(map_bytes), width=Cm(14))

    # -- Tile URL templates ordered by reliability ----
    _TILE_URLS = [
        "https://basemaps.cartocdn.com/light_all/{z}/{x}/{y}.png",
        "https://tile.openstreetmap.org/{z}/{x}/{y}.png",
    ]

    @staticmethod
    def _generate_static_map(
        lat: float, lon: float, zoom: int = 15, width: int = 600, height: int = 400
    ) -> bytes | None:
        """Generate a static map PNG using the staticmap library.

        Tries multiple tile sources for reliability.
        """
        try:
            from staticmap import StaticMap, CircleMarker
        except ImportError:
            logger.warning("staticmap not installed, skipping location map")
            return None

        for url_template in WordGenerator._TILE_URLS:
            try:
                m = StaticMap(
                    width, height,
                    url_template=url_template,
                    headers={"User-Agent": "SlopeDisasterReport/1.0"},
                    tile_request_timeout=10,
                    delay_between_retries=1,
                )
                marker = CircleMarker((lon, lat), "red", 12)
                m.add_marker(marker)
                image = m.render(zoom=zoom)
                buf = io.BytesIO()
                image.save(buf, format="PNG")
                buf.seek(0)
                return buf.getvalue()
            except Exception:
                logger.warning("Tile source %s failed, trying next", url_template)
                continue

        logger.error("All tile sources failed for location map")
        return None

    def _fill_photos(self, doc: DocxDocument, case: Case, manifest: EvidenceManifest | None) -> None:
        """Insert evidence photos after P21 (照片 section), then coordinates below.

        Photos are inserted consecutively with height=5.25cm, width=7cm.
        Coordinates are appended as text after the last photo.
        """
        paragraph = self._find_paragraph_containing(doc, "照片")
        if paragraph is None:
            return

        from docx.shared import Cm

        IMAGE_EXTENSIONS = {".jpg", ".jpeg", ".png", ".bmp", ".gif", ".webp", ".heic"}

        # Collect image evidence (by extension, not content_type which may be octet-stream)
        photos = []
        if manifest and manifest.evidence:
            photos = [
                ev for ev in manifest.evidence
                if Path(ev.evidence_path).suffix.lower() in IMAGE_EXTENSIONS
            ]
            photos.sort(key=lambda e: (e.photo_type or "Z", e.uploaded_at))

        # Insert photos
        last_p = paragraph
        if photos:
            img_p = self._insert_paragraph_after(paragraph, "")
            for ev in photos:
                photo_path = self._cases_dir / case.case_id / ev.evidence_path
                if not photo_path.exists():
                    logger.warning("Photo file not found: %s", photo_path)
                    continue
                run = img_p.add_run()
                try:
                    run.add_picture(str(photo_path), width=Cm(7), height=Cm(5.25))
                except Exception:
                    logger.exception("Failed to embed photo %s", photo_path.name)
                    continue
            last_p = img_p

        # Append coordinate text below photos (or below the 照片 heading if no photos)
        if case.primary_coordinate:
            coord_text = f"座標：N {case.primary_coordinate.lat:.4f}, E {case.primary_coordinate.lon:.4f}"
            coord_p = self._insert_paragraph_after(last_p, "")
            run = coord_p.add_run(coord_text)
            self._set_run_font(run)

    def _fill_analysis(self, doc: DocxDocument, case: Case, manifest: EvidenceManifest | None = None) -> None:
        """Fill analysis paragraph and toggle Table 1 damage mode checkboxes."""
        if case.analysis_review:
            paragraph = self._find_paragraph_containing(doc, "破壞模式與可能致災原因分析與檢討")
            if paragraph is not None:
                self._insert_paragraph_after(paragraph, case.analysis_review)

        # Table 1 checkbox toggling — photo annotation mapping by category
        if case.damage_mode_category in {"road_slope", "revetment_retaining", "bridge"}:
            self._fill_damage_mode_checkboxes(doc, case, manifest)

    @staticmethod
    def _toggle_checkbox_in_cell(cell: Any, target_text: str, check: bool) -> None:
        """Toggle a checkbox item in a table cell: □→■ (check) or ■→□ (uncheck).

        Handles cells containing multiple items separated by ' | '.
        Searches for target_text within the cell's paragraph runs.
        """
        old_char = "□" if check else "■"
        new_char = "■" if check else "□"
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                if target_text in run.text and old_char in run.text:
                    run.text = run.text.replace(old_char + target_text, new_char + target_text, 1)
                    return
            # Fallback: search in full paragraph text and rebuild
            full_text = paragraph.text
            if target_text in full_text and old_char in full_text:
                new_text = full_text.replace(old_char + target_text, new_char + target_text, 1)
                if new_text != full_text:
                    paragraph.clear()
                    run = paragraph.add_run(new_text)
                    WordGenerator._set_run_font(run)
                    return

    def _fill_damage_mode_checkboxes(self, doc: DocxDocument, case: Case, manifest: EvidenceManifest | None) -> None:
        """Toggle Table 1 (破壞模式與致災原因) checkboxes from photo annotations.

        Gate policy:
        1) Determine damage-mode checks by damage category.
        2) Check disaster causes only within matched damage-mode groups.
        """
        tags = self._extract_photo_tags(manifest)
        table = self._find_table_containing(doc, "破壞模式")
        if table is None:
            logger.warning("Table 1 not found when filling damage mode checkboxes")
            return

        def mark_damage(row_idx: int, label: str, should_check: bool) -> None:
            if not should_check or row_idx >= len(table.rows):
                return
            row = table.rows[row_idx]
            if len(row.cells) >= 2:
                self._check_table1_row(row.cells[1], label)

        def mark_cause(row_idx: int, label: str, should_check: bool) -> None:
            if not should_check or row_idx >= len(table.rows):
                return
            row = table.rows[row_idx]
            if len(row.cells) >= 4:
                self._check_table1_row(row.cells[3], label)

        p2_tags = tags.get("P2", {})
        p4_tags = tags.get("P4", {})
        p7_tags = tags.get("P7", {})

        if case.damage_mode_category == "road_slope":
            slope_locations = p2_tags.get("slope_location", set())
            geology_risks = p7_tags.get("geology_risk", set())
            slope_gradients = p4_tags.get("slope_gradient", set())

            mode_up = "cut_slope" in slope_locations
            mode_down = "fill_slope" in slope_locations
            mode_both = "both" in slope_locations

            mark_damage(12, "道路上方邊坡滑動", mode_up)
            mark_damage(13, "道路下方邊坡滑動", mode_down)
            mark_damage(14, "整體性破壞", mode_both)

            gate_road = mode_up or mode_down or mode_both
            steep_gradients = {"moderate", "steep", "very_steep", "cliff"}
            mark_cause(13, "土質鬆軟", gate_road and "soft_soil" in geology_risks)
            mark_cause(15, "坡度過大", gate_road and (bool(slope_gradients & steep_gradients) or "steep_slope" in geology_risks))
            return

        if case.damage_mode_category == "revetment_retaining":
            p1_tags = tags.get("P1", {})
            structure_locations = p2_tags.get("structure_location", set())
            visible_damage = p2_tags.get("visible_damage", set())
            site_risks = p1_tags.get("site_risks", set())
            water_body = p4_tags.get("water_body", set())
            drainage = p4_tags.get("drainage", set())
            water_signs = p4_tags.get("water_signs", set()) | p2_tags.get("water_signs", set())
            geology_risks = p7_tags.get("geology_risk", set())
            disaster_causes = p4_tags.get("disaster_cause", set())

            mode_revetment = bool(visible_damage & {"collapse", "displacement", "tilt", "scour", "undermining", "surface_erosion"})
            mode_river_structure = "riverside" in structure_locations and bool(visible_damage - {"no_visible_damage"})
            mode_upslope_retaining = "road_upslope" in structure_locations and bool(visible_damage - {"no_visible_damage"})
            mode_downslope_retaining = "road_downslope" in structure_locations and bool(visible_damage - {"no_visible_damage"})

            mark_damage(1, "護岸、擋土牆崩坍", mode_revetment)
            mark_damage(2, "河道內結構物破壞", mode_river_structure)
            mark_damage(10, "道路上方邊坡擋土牆破壞", mode_upslope_retaining)
            mark_damage(11, "道路下方邊坡擋土牆破壞", mode_downslope_retaining)

            gate_revetment = mode_revetment or mode_river_structure or mode_upslope_retaining or mode_downslope_retaining
            mark_cause(1, "水路流速過大，使基腳掏空或沖毀。", gate_revetment and bool(visible_damage & {"undermining", "foundation_exposed", "scour"}))
            mark_cause(2, "水路流速過大，護岸面被異物撞擊損毀。", gate_revetment and bool(visible_damage & {"crack", "spalling", "displacement", "tilt"}) and bool(water_body & {"river_high", "river_turbid", "bank_erosion"}))
            mark_cause(3, "水路流速過大，使護岸面被淘刷。", gate_revetment and bool(visible_damage & {"scour", "surface_erosion"}))
            mark_cause(4, "坡面無排水設施(自然邊坡)", gate_revetment and ("no_drainage" in drainage or "dry" in water_signs))
            mark_cause(5, "設計不足", gate_revetment and "design_inadequate" in disaster_causes)
            mark_cause(6, "坡面排水不良", gate_revetment and ("poor_drainage" in disaster_causes or "slope_poor_drain" in disaster_causes))
            mark_cause(7, "道路排水不良", gate_revetment and ("road_poor_drain" in disaster_causes or "subgrade_drain_poor" in disaster_causes))
            mark_cause(8, "存在介面", gate_revetment and "interface_exist" in geology_risks)
            mark_cause(9, "水路流量過大造成溢流，使臨水構造物損壞。", gate_revetment and bool(water_body & {"overtopping", "river_high"}))
            mark_cause(16, "路基缺口因道路設計排水不良", gate_revetment and "subgrade_gap" in site_risks and bool(drainage & {"blocked", "severe_blocked"}))
            mark_cause(17, "排水溝、集水井未定期清理所致", gate_revetment and bool(drainage & {"blocked", "severe_blocked", "catch_basin"}))
            return

        if case.damage_mode_category == "bridge":
            components = p2_tags.get("damaged_component", set())
            visible_damage = p2_tags.get("visible_damage", set())
            foundation_exposure = p2_tags.get("foundation_exposure", set())
            severity = p2_tags.get("severity", set())

            disaster_causes = p4_tags.get("disaster_cause", set())
            river_condition = p4_tags.get("river_condition", set())
            pier_scour = p4_tags.get("pier_scour", set())
            river_obstacles = p4_tags.get("river_obstacles", set())
            deck_overall = p4_tags.get("deck_overall", set())
            abutment_condition = p4_tags.get("abutment_condition", set())

            mode_18 = len(components & {"deck", "girder", "pier"}) >= 2 and bool(visible_damage & {"collapse", "displacement"}) and "critical" in severity
            mode_19 = bool(components & {"pier", "foundation"}) and ("foundation_exposed" in visible_damage or bool(foundation_exposure & {"depth_100_200", "depth_gt_200", "full_exposure"}))
            mode_20 = "pier" in components and "impact_damage" in visible_damage and bool(visible_damage & {"crack", "spalling"})
            mode_21 = "deck" in components and bool(visible_damage & {"tilt", "displacement"})
            mode_22 = bool(components & {"abutment", "wing_wall"}) and bool(visible_damage - {"no_visible_damage"})
            mode_23 = "girder" in components and "impact_damage" in visible_damage and bool(visible_damage & {"crack", "spalling", "rebar_exposed"})
            mode_24 = "girder" in components and "impact_damage" in visible_damage and bool(visible_damage & {"collapse", "displacement"})
            mode_25 = bool(components & {"railing", "parapet"}) and bool(visible_damage & {"crack", "spalling"})
            mode_26 = "drainage_blocked" in deck_overall and bool(river_obstacles & {"debris_pile", "waste"})

            mark_damage(18, "整跨落橋", mode_18)
            mark_damage(19, "橋墩基礎裸露", mode_19)
            mark_damage(20, "橋墩撞擊混凝土表面破裂", mode_20)
            mark_damage(21, "橋面傾斜", mode_21)
            mark_damage(22, "橋台翼牆破壞", mode_22)
            mark_damage(23, "橋梁大梁撞傷混凝土破裂鋼筋裸露", mode_23)
            mark_damage(24, "橋梁大梁撞斷混凝土破裂鋼筋裸露", mode_24)
            mark_damage(25, "橋面護欄破裂", mode_25)
            mark_damage(26, "橋面路燈損壞排水管阻塞", mode_26)

            gate_bridge = any([mode_18, mode_19, mode_20, mode_21, mode_22, mode_23, mode_24, mode_25, mode_26])
            mark_cause(18, "洪水沖刷，橋墩傾斜", gate_bridge and ("flood_scour" in disaster_causes or "flood" in river_condition) and ("tilt" in visible_damage or "pier" in components))
            mark_cause(19, "洪水沖刷掏空", gate_bridge and ("flood_scour" in disaster_causes or "undermined" in pier_scour or "scour" in visible_damage))
            mark_cause(20, "洪水夾雜石塊撞擊", gate_bridge and ("debris_impact" in disaster_causes or "impact_damage" in visible_damage))
            mark_cause(21, "基礎部份掏空，橋墩位移", gate_bridge and ("foundation_settlement" in disaster_causes or ("undermined" in pier_scour and "displacement" in visible_damage)))
            mark_cause(22, "洪水沖刷基礎橋台背牆位移", gate_bridge and "flood_scour" in disaster_causes and (bool(abutment_condition & {"backwall_damage", "wingwall_damage"}) or "displacement" in visible_damage))
            mark_cause(23, "翼牆與橋台旁防洪牆共構基礎遭洪水沖刷破壞", gate_bridge and "flood_scour" in disaster_causes and bool(components & {"wing_wall", "abutment"}))
            mark_cause(24, "洪水夾雜石塊，樹幹撞擊大梁", gate_bridge and "debris_impact" in disaster_causes and "girder" in components and "impact_damage" in visible_damage)
            mark_cause(25, "洪水淹沒橋面夾雜石塊撞擊", gate_bridge and ("flood" in river_condition or "high_water" in river_condition) and "deck" in components and "impact_damage" in visible_damage)
            mark_cause(26, "洪水夾雜樹枝、垃圾阻塞，路燈被颱風吹斷", gate_bridge and bool(river_obstacles & {"debris_pile", "waste"}) and "drainage_blocked" in deck_overall)

    @staticmethod
    def _check_table1_row(cell: Any, target_text: str) -> None:
        """Mark a 破壞模式/致災原因 row as checked by prepending ■ or replacing □→■.

        Table 1 rows have plain text (no pre-existing □/■). We prepend ■ to indicate checked.
        """
        for paragraph in cell.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
            # If target_text is found in the cell text
            if target_text in text:
                # Replace □ with ■ if checkbox prefix exists
                if "□" + target_text in text:
                    new_text = text.replace("□" + target_text, "■" + target_text, 1)
                    paragraph.clear()
                    run = paragraph.add_run(new_text)
                    WordGenerator._set_run_font(run)
                    return
                # Otherwise prepend ■ to indicate checked
                if not text.startswith("■"):
                    new_text = "■" + text
                    paragraph.clear()
                    run = paragraph.add_run(new_text)
                    WordGenerator._set_run_font(run)
                    return
    def _fill_cost(self, doc: DocxDocument, case: Case) -> None:
        """Fill cost breakdown table and total.

        Unit conversions:
        - case.estimated_cost is stored in 萬元 (ten-thousands NTD)
        - Word template P30 expects 仟元 (thousands NTD)
        - Conversion: 萬元 × 10 = 仟元
        """
        if case.cost_breakdown:
            paragraph = self._find_paragraph_containing(doc, "工程內容、數量及單價")
            if paragraph is not None:
                self._insert_cost_table(doc, paragraph, case.cost_breakdown)

        if case.estimated_cost is None:
            return

        # Convert 萬元 → 仟元 (× 10)
        cost_in_thousands = case.estimated_cost * 10
        total = self._format_number(cost_in_thousands)
        for paragraph in doc.paragraphs:
            stripped = paragraph.text.strip()
            if stripped.startswith("合計"):
                # Blank template has just "合計" (1 run); old template had "合計 ... 仟元" (2+ runs)
                if len(paragraph.runs) >= 2 and "仟元" in stripped:
                    paragraph.runs[0].text = f"合計 {total} "
                    paragraph.runs[1].text = "仟元"
                    self._set_run_font(paragraph.runs[0])
                    self._set_run_font(paragraph.runs[1])
                elif paragraph.runs:
                    paragraph.runs[0].text = f"合計 {total} 仟元"
                    self._set_run_font(paragraph.runs[0])
                break

    def _insert_cost_table(self, doc: DocxDocument, after_paragraph: Paragraph, items: list[CostBreakdownItem]) -> None:
        """Insert cost table after target paragraph.

        Unit conversions:
        - item.amount is stored in 元 (raw NTD)
        - Word template column header says 金額(仟元)
        - Conversion: 元 ÷ 1000 = 仟元
        """
        rows_data: list[list[str]] = []
        for item in items:
            if not item.item_name:
                continue
            # Convert amount from 元 to 仟元
            amount_in_thousands = item.amount / 1000 if item.amount else None
            rows_data.append(
                [
                    item.item_name,
                    item.unit or "-",
                    self._format_number(item.unit_price) or "-",
                    self._format_number(item.quantity) or "-",
                    self._format_number(amount_in_thousands) or "-",
                ]
            )

        if not rows_data:
            return

        table = doc.add_table(rows=1, cols=5)
        try:
            table.style = "Table Grid"
        except KeyError:
            logger.debug("Table Grid style not found in template; keeping default table style")

        headers = ["項目", "單位", "單價", "數量", "金額(仟元)"]
        for col_idx, header in enumerate(headers):
            header_cell = table.rows[0].cells[col_idx]
            self._set_paragraph_text(header_cell.paragraphs[0], header)

        for row_data in rows_data:
            row_cells = table.add_row().cells
            for col_idx, value in enumerate(row_data):
                self._set_paragraph_text(row_cells[col_idx].paragraphs[0], value)

        after_paragraph._element.addnext(table._element)

    def _fill_other_supplement(self, doc: DocxDocument, case: Case) -> None:
        """Fill other supplement text after section header."""
        if not case.other_supplement:
            return
        paragraph = self._find_paragraph_containing(doc, "其他補充事項")
        if paragraph is not None:
            self._insert_paragraph_after(paragraph, case.other_supplement)

    def _fill_soil_conservation(self, doc: DocxDocument, case: Case) -> None:
        """Fill soil conservation checkbox section."""
        if not case.soil_conservation:
            return

        start_idx = self._find_paragraph_index_containing(doc, "水土保持")
        if start_idx is None:
            return

        end_idx = self._find_paragraph_index_containing(doc, "是否位於國家公園範圍內")
        end_idx = end_idx if end_idx is not None else len(doc.paragraphs)
        paragraphs = doc.paragraphs[start_idx:end_idx]

        need_plan = case.soil_conservation.startswith("需要")
        no_plan = case.soil_conservation == "不需要"

        for paragraph in paragraphs:
            text = paragraph.text.strip()
            if text.startswith("1、是") and len(paragraph.runs) >= 2:
                paragraph.runs[1].text = ("■" if need_plan else "□") + "是。"
                self._set_run_font(paragraph.runs[1])
            elif text.startswith("2、否，請依下列檢核事項勾選") and len(paragraph.runs) >= 2:
                paragraph.runs[1].text = ("■" if no_plan else "□") + "否，請依下列檢核事項勾選："
                self._set_run_font(paragraph.runs[1])
            elif no_plan and text.startswith("(1)"):
                if paragraph.runs:
                    paragraph.runs[0].text = "■" + paragraph.runs[0].text
                    self._set_run_font(paragraph.runs[0])
                else:
                    self._set_paragraph_text(paragraph, "■" + text)
                break

    def _fill_national_park(self, doc: DocxDocument, case: Case) -> None:
        """Fill national park checkbox section."""
        start_idx = self._find_paragraph_index_containing(doc, "是否位於國家公園範圍內")
        if start_idx is None:
            return
        paragraphs = doc.paragraphs[start_idx:]
        in_park = bool(case.national_park)

        yes_subparagraphs: list[Paragraph] = []
        no_subparagraphs: list[Paragraph] = []

        for paragraph in paragraphs:
            text = paragraph.text.strip()
            if text.startswith("1、是，請依下列檢核事項勾選") and len(paragraph.runs) >= 2:
                paragraph.runs[1].text = ("■" if in_park else "□") + "是，請依下列檢核事項勾選："
                self._set_run_font(paragraph.runs[1])
            elif text == "2、否。" and len(paragraph.runs) >= 2:
                paragraph.runs[1].text = ("■" if not in_park else "□") + "否。"
                self._set_run_font(paragraph.runs[1])
            elif text == "是":
                yes_subparagraphs.append(paragraph)
            elif text.startswith("否，"):
                no_subparagraphs.append(paragraph)

        for paragraph in yes_subparagraphs[:2]:
            self._set_paragraph_text(paragraph, ("■" if in_park else "□") + "是")

        for paragraph in no_subparagraphs[:2]:
            original_text = paragraph.text
            self._set_paragraph_text(paragraph, ("□" if in_park else "■") + original_text)

    def _fill_hazard_table(self, doc: DocxDocument, case: Case, manifest: EvidenceManifest | None = None) -> None:
        """Fill Table 2 (工址環境調查及危害辨識) with ■/□ checkboxes and 備註 notes.

        Two-phase approach:
        1. Toggle ■/□ in col 1 (現地狀況) and col 2 (工址風險) based on photo tag 勾稽
        2. Write 備註 notes to col 3 based on site_survey selections (existing behavior)
        """
        if len(doc.tables) < 3:
            return

        table = doc.tables[2]
        tags = self._extract_photo_tags(manifest)

        # --- Phase 1: ■/□ 勾稽 in cols 1-2 ---
        self._toggle_hazard_checkboxes(table, tags, case)

        # --- Phase 2: 備註 notes in col 3 (existing logic) ---
        if not case.site_survey:
            return

        checked_items = {item.item_id: item for item in case.site_survey if item.checked}
        if not checked_items:
            return

        row_item_map: dict[int, list[str]] = {
            1: ["upslope_rockfall"],
            2: ["upslope_collapse"],
            3: ["downslope_subgrade_gap", "downslope_settlement"],
            4: [
                "structure_rebar_exposed",
                "structure_dangerous_tree",
                "structure_guardrail",
                "structure_pothole",
                "structure_utility_pole",
            ],
            5: ["bridge_river_adjacent", "bridge_river_meander_erosion"],
            6: ["other_weather"],
            7: ["other_snow"],
        }

        for row_idx, item_ids in row_item_map.items():
            if row_idx >= len(table.rows):
                continue
            notes: list[str] = []
            for item_id in item_ids:
                item = checked_items.get(item_id)
                if item is None:
                    continue
                note_text = f"✓ {item.item_name or item_id}"
                if item.note:
                    note_text += f"（{item.note}）"
                notes.append(note_text)

            if notes:
                row = table.rows[row_idx]
                if len(row.cells) > 3:
                    self._set_cell_text(row.cells[3], "\n".join(notes))

        if case.hazard_supplement:
            target_row_idx = min(7, len(table.rows) - 1)
            row = table.rows[target_row_idx]
            if len(row.cells) <= 3:
                return
            cell = row.cells[3]
            existing = cell.text.strip()
            supplement = f"補充：{case.hazard_supplement}"
            merged = f"{existing}\n{supplement}" if existing else supplement
            self._set_cell_text(cell, merged)

    def _toggle_hazard_checkboxes(self, table: Any, tags: dict[str, dict[str, set[str]]], case: Case) -> None:
        """Toggle ■/□ in Table 2 cols 1 (現地狀況) and col 2 (工址風險) based on photo tag 勾稽.

        Cell content uses ' | ' to separate sub-items. Each sub-item can be toggled independently.
        Template uses □ for unchecked, we replace with ■ for checked.
        """
        p1_tags = tags.get("P1", {})
        p2_tags = tags.get("P2", {})

        site_risks = p1_tags.get("site_risks", set())
        visible_damage = p2_tags.get("visible_damage", set())

        # ---- Row 1: 上邊坡 - 落石甫清理 ----
        # 現地狀況 col 1: 落石甫清理 (always matches as row identifier)
        # 工址風險 col 2: 物體飛落危害 | 物體倒塌/崩塌
        r1_collapse = (
            "upslope_rockfall" in site_risks
        )
        r1_flying = (
            "hanging_rock" in visible_damage
            or "isolated_rock_pile" in visible_damage
        )
        if r1_collapse or r1_flying:
            self._toggle_hazard_row(table, 1, {
                1: [],  # col 1 現地狀況: no sub-items to toggle
                2: (["物體倒塌/崩塌"] if r1_collapse else [])
                   + (["物體飛落危害"] if r1_flying else []),
            })

        # ---- Row 2: 上邊坡 - 局部崩塌、土石流 ----
        # 工址風險 col 2: 物體飛落危害 | 物體倒塌/崩塌
        r2_collapse = (
            "debris_avalanche" in visible_damage
            or "rock_mass_slide" in visible_damage
            or "debris_flow" in visible_damage
            or "collapse_sign" in site_risks
            or "debris_flow_sign" in site_risks
        )
        r2_flying = r2_collapse and "slope_debris_deposit" in visible_damage
        if r2_collapse or r2_flying:
            self._toggle_hazard_row(table, 2, {
                1: [],  # col 1: no sub-items to toggle
                2: (["物體倒塌/崩塌"] if r2_collapse else [])
                   + (["物體飛落危害"] if r2_flying else []),
            })

        # ---- Row 3: 下邊坡 ----
        # 現地狀況 col 1: 路基缺口 | 路基下陷
        # 工址風險 col 2: 墜落、滾落危害 | 衝撞、被撞危害
        r3_gap = "subgrade_gap" in site_risks
        r3_sink = "subsidence" in site_risks
        if r3_gap or r3_sink:
            self._toggle_hazard_row(table, 3, {
                1: (["路基缺口"] if r3_gap else [])
                   + (["路基下陷"] if r3_sink else []),
                2: ["墜落、滾落危害", "衝撞、被撞危害"],  # both checked if any row 3 trigger
            })

        # ---- Row 4: 結構物、路側及路面 ----
        # 現地狀況 col 1: 結構物鋼筋裸露 | 危木倒塌 | 護欄損壞 | 路面坑洞 | 路側電桁倒塌
        # 工址風險 col 2: 被刺、割、擦傷 | 物體倒塌/崩塌 | 墜落、滾落危害 | 跌倒 | 感電危害
        r4_rebar = "rebar_exposed" in site_risks
        r4_tree = "hazard_tree" in site_risks
        r4_guard = "guardrail_damage" in site_risks
        r4_pot = "pothole" in site_risks
        r4_pole = "utility_pole_tilt" in site_risks

        r4_col1 = []
        r4_col2 = []
        if r4_rebar:
            r4_col1.append("結構物鋼筋裸露")
            r4_col2.append("被刺、割、擦傷")
        if r4_tree:
            r4_col1.append("危木倒塌")
            r4_col2.append("物體倒塌/崩塌")
        if r4_guard:
            r4_col1.append("護欄損壞")
            r4_col2.append("墜落、滾落危害")
        if r4_pot:
            r4_col1.append("路面坑洞")
            r4_col2.append("跌倒")
        if r4_pole:
            r4_col1.append("路側電桁倒塌")
            r4_col2.append("感電危害")
        if r4_col1 or r4_col2:
            self._toggle_hazard_row(table, 4, {1: r4_col1, 2: r4_col2})

        # ---- Row 5: 橋樑、河道 ----
        # 現地狀況 col 1: 臨河作業 | 曲流攻擊面
        # 工址風險 col 2: 墜落危害 | 溺斃危害
        r5_river = "riverside_work" in site_risks
        r5_meander = "meander_attack" in site_risks

        r5_col1 = []
        r5_col2 = []
        if r5_river:
            r5_col1.append("臨河作業")
            r5_col2.extend(["墜落危害", "溺斃危害"])
        if r5_meander:
            r5_col1.append("曲流攻擊面")
            if "墜落危害" not in r5_col2:
                r5_col2.append("墜落危害")
            if "溺斃危害" not in r5_col2:
                r5_col2.append("溺斃危害")
        if r5_col1 or r5_col2:
            self._toggle_hazard_row(table, 5, {1: r5_col1, 2: r5_col2})

        # Rows 6-7 (其他 天氣/降雪): 不自動勾稽，保留手動

        logger.info(
            "勾稽 Table 2: site_risks=%s, visible_damage=%s",
            site_risks, visible_damage,
        )

    @staticmethod
    def _toggle_hazard_row(table: Any, row_idx: int, col_items: dict[int, list[str]]) -> None:
        """Toggle ■/□ for specific sub-items in a Table 2 row.

        Args:
            table: The docx table object
            row_idx: Row index in table
            col_items: {col_idx: [list of sub-item texts to check ■]}
        """
        if row_idx >= len(table.rows):
            return
        row = table.rows[row_idx]
        for col_idx, items_to_check in col_items.items():
            if col_idx >= len(row.cells) or not items_to_check:
                continue
            cell = row.cells[col_idx]
            for paragraph in cell.paragraphs:
                text = paragraph.text
                changed = False
                for item_text in items_to_check:
                    # Replace □+item with ■+item
                    target_unchecked = "□" + item_text
                    target_checked = "■" + item_text
                    if target_unchecked in text:
                        text = text.replace(target_unchecked, target_checked, 1)
                        changed = True
                    elif item_text in text and target_checked not in text:
                        # Item exists without □/■ prefix — prepend ■
                        text = text.replace(item_text, "■" + item_text, 1)
                        changed = True
                if changed:
                    paragraph.clear()
                    run = paragraph.add_run(text)
                    WordGenerator._set_run_font(run)

    @classmethod
    def _set_cell_text(cls, cell: Any, text: str) -> None:
        """Replace first paragraph text in a cell and apply font."""
        if not cell.paragraphs:
            from docx.oxml.ns import qn
            cell._element.append(cell._element.makeelement(qn('w:p'), {}))
        paragraph = cell.paragraphs[0]
        paragraph.clear()
        for idx, part in enumerate(text.split("\n")):
            if idx > 0:
                paragraph.add_run().add_break()
            run = paragraph.add_run(part)
            cls._set_run_font(run)

    def _fill_signatures(self, doc: DocxDocument, case: Case) -> None:
        """Fill signature table and reporting agency line."""
        if len(doc.tables) >= 4 and case.created_by and case.created_by.real_name:
            signature_cell = doc.tables[3].rows[1].cells[0]
            self._set_cell_text(signature_cell, case.created_by.real_name)

        if case.reporting_agency:
            paragraph = self._find_paragraph_containing(doc, "提報機關")
            if paragraph is not None:
                if paragraph.runs:
                    paragraph.runs[0].text = f"提報機關：{case.reporting_agency}"
                    self._set_run_font(paragraph.runs[0])
                else:
                    self._set_paragraph_text(paragraph, f"提報機關：{case.reporting_agency}")
