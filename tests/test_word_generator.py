from __future__ import annotations

import io
import sys
import types
from pathlib import Path
from unittest.mock import MagicMock, patch

import pytest
from docx import Document
from docx.document import Document as DocxDocument
from PIL import Image

from app.models.case import Case, CoordinateCandidate, CostBreakdownItem, CreatedBy, SiteSurveyItem
from app.models.evidence import AnnotationTag, EvidenceManifest, EvidenceMetadata, PhotoAnnotations
from app.services.word_generator import WordGenerator


@pytest.fixture
def template_path() -> Path:
    project_root = Path(__file__).resolve().parents[1]
    return project_root / "Input" / "公路災害工程內容概述表_空白.docx"


@pytest.fixture
def generator(template_path: Path, tmp_path: Path) -> WordGenerator:
    return WordGenerator(template_path=template_path, cases_dir=tmp_path)


def _document_text(doc: DocxDocument) -> tuple[str, str]:
    all_paragraph_text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    all_table_text = "\n".join(cell.text for table in doc.tables for row in table.rows for cell in row.cells)
    return all_paragraph_text, all_table_text


def _section_text(doc: DocxDocument, start_marker: str, end_marker: str | None = None) -> str:
    start_idx = next((i for i, p in enumerate(doc.paragraphs) if start_marker in p.text), None)
    if start_idx is None:
        return ""
    end_idx = len(doc.paragraphs)
    if end_marker is not None:
        match_idx = next(
            (i for i, p in enumerate(doc.paragraphs[start_idx + 1 :], start=start_idx + 1) if end_marker in p.text),
            None,
        )
        if match_idx is not None:
            end_idx = match_idx
    return "\n".join(p.text for p in doc.paragraphs[start_idx:end_idx])


def _find_table_containing(doc: DocxDocument, text: str):
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if text in cell.text:
                    return table
    raise AssertionError(f"Table containing '{text}' not found")


def test_generate_smoke_fills_core_fields(generator: WordGenerator) -> None:
    case = Case(
        case_id="case_20260301_0001",
        reporting_year="115",
        disaster_type="一般",
        processing_type="搶修",
        project_name="測試工程",
        disaster_date="2025-08-15",
        town_name="復興",
        village_name="高義",
        nearby_landmark="台7線北側邊坡",
        repeat_disaster="是",
        repeat_disaster_year="98",
        description="災損說明測試段落",
        original_protection="原設計保護型式測試段落",
        analysis_review="分析與檢討測試段落",
        cost_breakdown=[
            CostBreakdownItem(
                item_id="1",
                item_name="擋土牆修復",
                unit="式",
                unit_price=1200000,
                quantity=2.0,
                amount=2400000,
            )
        ],
        estimated_cost=240.0,
        other_supplement="其他補充測試段落",
        soil_conservation="需要已核定",
        national_park="太魯閣國家公園",
        reporting_agency="交通部公路局北區養護工程分局",
        primary_coordinate=CoordinateCandidate(lat=24.123456, lon=121.654321, source="manual", confidence=1.0),
        site_survey=[
            SiteSurveyItem(
                category_id="upslope",
                item_id="upslope_rockfall",
                item_name="落石待清理",
                checked=True,
                note="有新生落石",
            )
        ],
        created_by=CreatedBy(user_id="u001", real_name="王小明"),
    )

    generated = generator.generate(case=case, manifest=EvidenceManifest(case_id=case.case_id))

    assert isinstance(generated, bytes)
    assert generated.startswith(b"PK")

    doc = Document(io.BytesIO(generated))
    all_paragraph_text, all_table_text = _document_text(doc)

    assert "測試工程" in all_paragraph_text
    assert "災損說明測試段落" in all_paragraph_text
    assert "原設計保護型式測試段落" in all_paragraph_text
    assert "分析與檢討測試段落" in all_paragraph_text
    assert "其他補充測試段落" in all_paragraph_text
    assert "■一般災害" in all_paragraph_text
    assert "■搶修" in all_paragraph_text
    assert "■是" in all_paragraph_text
    assert "121.6543" in all_paragraph_text  # coordinate text below photos
    assert "24.1235" in all_paragraph_text  # coordinate text below photos
    assert "落石待清理" in all_table_text
    assert "王小明" in all_table_text


def test_generate_basic_fields_project_name(generator: WordGenerator) -> None:
    case = Case(case_id="case_basic_project", project_name="北橫搶修工程")

    generated = generator.generate(case=case)
    doc = Document(io.BytesIO(generated))

    assert "工程名稱：北橫搶修工程" in doc.paragraphs[4].text


def test_generate_date_fill_iso(generator: WordGenerator) -> None:
    case = Case(case_id="case_date_iso", disaster_date="2025-08-15")

    generated = generator.generate(case=case)
    doc = Document(io.BytesIO(generated))
    date_text = doc.paragraphs[5].text

    assert "114" in date_text
    assert "8" in date_text
    assert "15" in date_text


def test_generate_date_fill_minguo(generator: WordGenerator) -> None:
    case = Case(case_id="case_date_minguo", disaster_date="114年8月15日")

    generated = generator.generate(case=case)
    doc = Document(io.BytesIO(generated))
    date_text = doc.paragraphs[5].text

    assert "114" in date_text
    assert "8" in date_text
    assert "15" in date_text


def test_generate_date_fill_minguo_slash(generator: WordGenerator) -> None:
    case = Case(case_id="case_date_minguo_slash", disaster_date="114/03/01")

    generated = generator.generate(case=case)
    doc = Document(io.BytesIO(generated))
    date_text = doc.paragraphs[5].text

    assert "114" in date_text
    assert "3" in date_text
    assert "1" in date_text


def test_generate_with_coordinates(generator: WordGenerator) -> None:
    case = Case(
        case_id="case_coordinates",
        primary_coordinate=CoordinateCandidate(lat=23.987654, lon=121.123456, source="manual", confidence=1.0),
    )

    generated = generator.generate(case=case)
    doc = Document(io.BytesIO(generated))
    all_paragraph_text = "\n".join(p.text for p in doc.paragraphs)

    # Coordinates are now appended as text below the photos section
    assert "121.1235" in all_paragraph_text
    assert "23.9877" in all_paragraph_text


def test_generate_with_cost_breakdown(generator: WordGenerator) -> None:
    case = Case(
        case_id="case_cost",
        cost_breakdown=[
            CostBreakdownItem(
                item_id="1",
                item_name="鋼軌護欄修復",
                unit="m",
                unit_price=3500,
                quantity=10,
                amount=35000,
            )
        ],
        estimated_cost=3.5,
    )

    generated = generator.generate(case=case)
    doc = Document(io.BytesIO(generated))

    table_text = "\n".join(cell.text for table in doc.tables for row in table.rows for cell in row.cells)
    paragraph_text = "\n".join(p.text for p in doc.paragraphs)

    assert "金額(仟元)" in table_text
    assert "鋼軌護欄修復" in table_text
    assert "合計 35" in paragraph_text


def test_generate_with_soil_conservation_yes(generator: WordGenerator) -> None:
    case = Case(case_id="case_soil_yes", soil_conservation="需要已核定")

    generated = generator.generate(case=case)
    doc = Document(io.BytesIO(generated))
    soil_text = _section_text(doc, "水土保持", "是否位於國家公園範圍內")

    assert "■是。" in soil_text
    assert "□否，請依下列檢核事項勾選：" in soil_text


def test_generate_with_soil_conservation_no(generator: WordGenerator) -> None:
    case = Case(case_id="case_soil_no", soil_conservation="不需要")

    generated = generator.generate(case=case)
    doc = Document(io.BytesIO(generated))
    soil_text = _section_text(doc, "水土保持", "是否位於國家公園範圍內")

    assert "□是。" in soil_text
    assert "■否，請依下列檢核事項勾選：" in soil_text


def test_generate_with_national_park(generator: WordGenerator) -> None:
    case = Case(case_id="case_park_yes", national_park="太魯閣國家公園")

    generated = generator.generate(case=case)
    doc = Document(io.BytesIO(generated))
    park_text = _section_text(doc, "是否位於國家公園範圍內")

    assert "■是，請依下列檢核事項勾選：" in park_text
    assert "□否。" in park_text


def test_generate_without_national_park(generator: WordGenerator) -> None:
    case = Case(case_id="case_park_no", national_park="")

    generated = generator.generate(case=case)
    doc = Document(io.BytesIO(generated))
    park_text = _section_text(doc, "是否位於國家公園範圍內")

    assert "□是，請依下列檢核事項勾選：" in park_text
    assert "■否。" in park_text


def test_generate_with_site_survey(generator: WordGenerator) -> None:
    case = Case(
        case_id="case_site_survey",
        cost_breakdown=[
            CostBreakdownItem(
                item_id="1",
                item_name="臨時排水改善",
                unit="式",
                unit_price=1.0,
                quantity=1.0,
                amount=1.0,
            )
        ],
        site_survey=[
            SiteSurveyItem(
                category_id="upslope",
                item_id="upslope_rockfall",
                item_name="落石待清理",
                checked=True,
                note="有新生落石",
            )
        ],
    )

    generated = generator.generate(case=case)
    doc = Document(io.BytesIO(generated))

    all_table_text = "\n".join(cell.text for table in doc.tables for row in table.rows for cell in row.cells)
    assert "✓ 落石待清理（有新生落石）" in all_table_text


def test_hazard_checkbox_uses_photo_tags_not_site_survey(generator: WordGenerator) -> None:
    case = Case(
        case_id="case_hazard_checkbox_no_survey_fallback",
        cost_breakdown=[
            CostBreakdownItem(item_id="1", item_name="測試工項", unit="式", unit_price=1, quantity=1, amount=1)
        ],
        site_survey=[
            SiteSurveyItem(
                category_id="upslope",
                item_id="upslope_rockfall",
                item_name="落石待清理",
                checked=True,
            )
        ],
    )

    generated = generator.generate(case=case, manifest=EvidenceManifest(case_id=case.case_id, evidence=[]))
    doc = Document(io.BytesIO(generated))
    row1_col2 = doc.tables[2].rows[1].cells[2].text

    assert "■" not in row1_col2


def test_hazard_checkbox_checks_when_photo_tags_match(generator: WordGenerator) -> None:
    case = Case(
        case_id="case_hazard_checkbox_from_tags",
        cost_breakdown=[
            CostBreakdownItem(item_id="1", item_name="測試工項", unit="式", unit_price=1, quantity=1, amount=1)
        ],
    )
    manifest = EvidenceManifest(
        case_id=case.case_id,
        evidence=[
            EvidenceMetadata(
                evidence_id="ev_001",
                sha256="abc123",
                original_filename="p1.jpg",
                content_type="image/jpeg",
                photo_type="P1",
                evidence_path="evidence/abc123.jpg",
                annotations=PhotoAnnotations(
                    tags=[AnnotationTag(category="site_risks", tag_id="upslope_rockfall", label="上邊坡落石")]
                ),
            )
        ],
    )

    generated = generator.generate(case=case, manifest=manifest)
    doc = Document(io.BytesIO(generated))
    row1_col2 = doc.tables[2].rows[1].cells[2].text

    assert "■" in row1_col2


def test_damage_mode_revetment_mapping_with_cause_gate(generator: WordGenerator) -> None:
    case = Case(case_id="case_damage_revetment", damage_mode_category="revetment_retaining")
    manifest = EvidenceManifest(
        case_id=case.case_id,
        evidence=[
            EvidenceMetadata(
                evidence_id="ev_p2",
                sha256="sha_p2",
                original_filename="p2.jpg",
                content_type="image/jpeg",
                photo_type="P2",
                evidence_path="evidence/sha_p2.jpg",
                annotations=PhotoAnnotations(
                    tags=[
                        AnnotationTag(category="structure_location", tag_id="riverside", label="臨河側"),
                        AnnotationTag(category="visible_damage", tag_id="collapse", label="崩塌"),
                        AnnotationTag(category="visible_damage", tag_id="undermining", label="掏空"),
                    ]
                ),
            ),
            EvidenceMetadata(
                evidence_id="ev_p4",
                sha256="sha_p4",
                original_filename="p4.jpg",
                content_type="image/jpeg",
                photo_type="P4",
                evidence_path="evidence/sha_p4.jpg",
                annotations=PhotoAnnotations(
                    tags=[
                        AnnotationTag(category="water_body", tag_id="river_high", label="河水暴漲"),
                        AnnotationTag(category="water_body", tag_id="overtopping", label="溢流"),
                    ]
                ),
            ),
            EvidenceMetadata(
                evidence_id="ev_p7",
                sha256="sha_p7",
                original_filename="p7.jpg",
                content_type="image/jpeg",
                photo_type="P7",
                evidence_path="evidence/sha_p7.jpg",
                annotations=PhotoAnnotations(
                    tags=[AnnotationTag(category="geology_risk", tag_id="interface_exist", label="存在介面")]
                ),
            ),
        ],
    )

    generated = generator.generate(case=case, manifest=manifest)
    doc = Document(io.BytesIO(generated))
    damage_table = _find_table_containing(doc, "破壞模式")
    all_text = "\n".join(cell.text for row in damage_table.rows for cell in row.cells)

    assert "■護岸、擋土牆崩坍" in all_text
    assert "■河道內結構物破壞" in all_text
    assert "■水路流速過大，使基腳掏空或沖毀。" in all_text
    assert "■存在介面" in all_text


def test_damage_mode_revetment_covers_retaining_wall_and_drainage_causes(generator: WordGenerator) -> None:
    case = Case(case_id="case_damage_revetment_full", damage_mode_category="revetment_retaining")
    manifest = EvidenceManifest(
        case_id=case.case_id,
        evidence=[
            EvidenceMetadata(
                evidence_id="ev_p1",
                sha256="sha_p1",
                original_filename="p1.jpg",
                content_type="image/jpeg",
                photo_type="P1",
                evidence_path="evidence/sha_p1.jpg",
                annotations=PhotoAnnotations(
                    tags=[AnnotationTag(category="site_risks", tag_id="subgrade_gap", label="路基缺口")]
                ),
            ),
            EvidenceMetadata(
                evidence_id="ev_p2",
                sha256="sha_p2x",
                original_filename="p2x.jpg",
                content_type="image/jpeg",
                photo_type="P2",
                evidence_path="evidence/sha_p2x.jpg",
                annotations=PhotoAnnotations(
                    tags=[
                        AnnotationTag(category="structure_location", tag_id="road_upslope", label="道路上方"),
                        AnnotationTag(category="structure_location", tag_id="road_downslope", label="道路下方"),
                        AnnotationTag(category="visible_damage", tag_id="crack", label="裂縫"),
                    ]
                ),
            ),
            EvidenceMetadata(
                evidence_id="ev_p4",
                sha256="sha_p4x",
                original_filename="p4x.jpg",
                content_type="image/jpeg",
                photo_type="P4",
                evidence_path="evidence/sha_p4x.jpg",
                annotations=PhotoAnnotations(
                    tags=[AnnotationTag(category="drainage", tag_id="severe_blocked", label="嚴重堵塞")]
                ),
            ),
        ],
    )

    generated = generator.generate(case=case, manifest=manifest)
    doc = Document(io.BytesIO(generated))
    damage_table = _find_table_containing(doc, "破壞模式")
    all_text = "\n".join(cell.text for row in damage_table.rows for cell in row.cells)

    assert "■道路上方邊坡擋土牆破壞" in all_text
    assert "■道路下方邊坡擋土牆破壞" in all_text
    assert "■路基缺口因道路設計排水不良" in all_text
    assert "■排水溝、集水井未定期清理所致" in all_text


def test_damage_mode_bridge_mapping_with_cause_gate(generator: WordGenerator) -> None:
    case = Case(case_id="case_damage_bridge", damage_mode_category="bridge")
    manifest = EvidenceManifest(
        case_id=case.case_id,
        evidence=[
            EvidenceMetadata(
                evidence_id="ev_p2_bridge",
                sha256="sha_p2_bridge",
                original_filename="p2_bridge.jpg",
                content_type="image/jpeg",
                photo_type="P2",
                evidence_path="evidence/sha_p2_bridge.jpg",
                annotations=PhotoAnnotations(
                    tags=[
                        AnnotationTag(category="damaged_component", tag_id="pier", label="橋墩"),
                        AnnotationTag(category="visible_damage", tag_id="impact_damage", label="撞擊損壞"),
                        AnnotationTag(category="visible_damage", tag_id="crack", label="裂縫"),
                    ]
                ),
            ),
            EvidenceMetadata(
                evidence_id="ev_p4_bridge",
                sha256="sha_p4_bridge",
                original_filename="p4_bridge.jpg",
                content_type="image/jpeg",
                photo_type="P4",
                evidence_path="evidence/sha_p4_bridge.jpg",
                annotations=PhotoAnnotations(
                    tags=[AnnotationTag(category="disaster_cause", tag_id="debris_impact", label="漂流物撞擊")]
                ),
            ),
        ],
    )

    generated = generator.generate(case=case, manifest=manifest)
    doc = Document(io.BytesIO(generated))
    damage_table = _find_table_containing(doc, "破壞模式")
    all_text = "\n".join(cell.text for row in damage_table.rows for cell in row.cells)

    assert "■橋墩撞擊混凝土表面破裂" in all_text
    assert "■洪水夾雜石塊撞擊" in all_text


def test_generate_with_photos(generator: WordGenerator, tmp_path: Path) -> None:
    case_id = "case_photos"
    case = Case(case_id=case_id)

    temp_jpg = tmp_path / "temp_test.jpg"
    Image.new("RGB", (100, 100), "red").save(temp_jpg, "JPEG")

    evidence_dir = tmp_path / case_id / "evidence"
    evidence_dir.mkdir(parents=True, exist_ok=True)
    photo_path = evidence_dir / "abc123.jpg"
    _ = photo_path.write_bytes(temp_jpg.read_bytes())

    manifest = EvidenceManifest(
        case_id=case_id,
        evidence=[
            EvidenceMetadata(
                evidence_id="ev_001",
                sha256="abc123",
                original_filename="test.jpg",
                content_type="image/jpeg",
                photo_type="P1",
                photo_type_name="遠景照",
                evidence_path="evidence/abc123.jpg",
            )
        ],
    )

    generated = generator.generate(case=case, manifest=manifest)
    doc = Document(io.BytesIO(generated))
    all_paragraph_text, _ = _document_text(doc)

    # Photos are now inserted without captions (consecutively)
    assert len(doc.inline_shapes) >= 1


def test_generate_with_location_map(generator: WordGenerator) -> None:
    case = Case(
        case_id="case_map",
        primary_coordinate=CoordinateCandidate(lat=24.111111, lon=121.222222, source="manual", confidence=1.0),
    )

    fake_staticmap_module = types.ModuleType("staticmap")
    setattr(fake_staticmap_module, "StaticMap", object)
    setattr(fake_staticmap_module, "CircleMarker", object)

    class _FakeMap:
        def __init__(self) -> None:
            self.marker: object | None = None
            self.zoom: int | None = None

        def add_marker(self, marker: object) -> None:
            self.marker = marker

        def render(self, *, zoom: int) -> Image.Image:
            self.zoom = zoom
            return Image.new("RGB", (320, 180), "blue")

    with patch.dict(sys.modules, {"staticmap": fake_staticmap_module}):
        with patch("staticmap.StaticMap") as static_map_cls, patch("staticmap.CircleMarker") as circle_marker_cls:
            fake_map = _FakeMap()
            static_map_cls.return_value = fake_map
            circle_marker = MagicMock()
            circle_marker_cls.return_value = circle_marker

            generated = generator.generate(case=case)

    doc = Document(io.BytesIO(generated))

    static_map_cls.assert_called_once()  # params include url_template, headers, etc.
    circle_marker_cls.assert_called_once_with((121.222222, 24.111111), "red", 12)
    assert fake_map.marker is circle_marker
    assert fake_map.zoom == 15
    assert len(doc.inline_shapes) >= 1


def test_generate_returns_valid_docx_bytes(generator: WordGenerator) -> None:
    case = Case(case_id="case_magic_bytes")

    generated = generator.generate(case=case)

    assert isinstance(generated, bytes)
    assert generated.startswith(b"PK")


def test_template_not_found_raises(tmp_path: Path) -> None:
    with pytest.raises(FileNotFoundError):
        _ = WordGenerator(template_path=tmp_path / "missing_template.docx", cases_dir=tmp_path)


def test_generate_empty_case(generator: WordGenerator) -> None:
    case = Case(case_id="case_empty")

    generated = generator.generate(case=case)
    doc = Document(io.BytesIO(generated))

    assert generated.startswith(b"PK")
    assert len(doc.paragraphs) > 0


def test_completeness_full_case(generator: WordGenerator) -> None:
    """完整填寫的 Case 應該有 100% 完成度。"""
    case = Case(
        case_id="case_complete",
        reporting_year="114",
        disaster_type="一般",
        processing_type="搶修",
        project_name="測試工程",
        disaster_date="2025-08-15",
        town_name="復興",
        village_name="高義",
        nearby_landmark="台7線北側邊坡",
        repeat_disaster="是",
        primary_coordinate=CoordinateCandidate(lat=24.1, lon=121.6, source="manual", confidence=1.0),
        damage_mode_name="岩石崩塌",
        damage_cause_names=["豪雨"],
        description="測試描述",
        photo_count=2,
        original_protection="重力式擋土牆",
        analysis_review="分析內容",
        estimated_cost=240.0,
        cost_breakdown=[CostBreakdownItem(item_id="1", item_name="修復", amount=2400000)],
        design_doc_evidence_id="ev_001",
        soil_conservation="不需要",
        safety_assessment="安全",
        site_survey=[SiteSurveyItem(category_id="upslope", item_id="test", item_name="test", checked=True)],
        other_supplement="補充",
        reporting_agency="交通部公路局北區養護工程分局",
        created_by=CreatedBy(user_id="u001", real_name="王小明"),
    )
    result = WordGenerator.calculate_completeness(case)
    assert result["filled"] == 25
    assert result["total"] == 25
    assert result["percentage"] == 100
    assert result["missing"] == []


def test_completeness_minimal_case(generator: WordGenerator) -> None:
    """空 Case 應有低完成度，且列出缺填必填欄位。"""
    case = Case(case_id="case_empty")
    result = WordGenerator.calculate_completeness(case)
    assert result["percentage"] < 50
    assert result["total"] == 25
    assert len(result["missing"]) > 0
    missing_names = [m["name"] for m in result["missing"]]
    assert "工程名稱" in missing_names
    assert "災害日期" in missing_names


def test_completeness_partial_case(generator: WordGenerator) -> None:
    """部分填寫的 Case 完成度介於 0~100%，且 missing 只列未填。"""
    case = Case(
        case_id="case_partial",
        reporting_year="114",
        disaster_type="一般",
        processing_type="搶修",
        project_name="工程A",
        disaster_date="2025-01-01",
        reporting_agency="交通部公路局北區養護工程分局",
        created_by=CreatedBy(user_id="u001", real_name="王小明"),
    )
    result = WordGenerator.calculate_completeness(case)
    assert 0 < result["percentage"] < 100
    assert result["filled"] == result["total"] - len(result["missing"])
