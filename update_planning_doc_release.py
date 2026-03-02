from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt


DOC_PATH = Path(
    r"C:\Users\Sad Kevin\Desktop\邊坡災害通報與資訊整合管理系統\邊坡災害通報與資訊整合管理系統_規劃文件.docx"
)
TITLE = "第14章 定版功能與自動報告勾稽原則"


def _set_run_font(run) -> None:
    run.font.name = "微軟正黑體"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微軟正黑體")
    run.font.size = Pt(12)


def _set_paragraph_font(paragraph) -> None:
    if not paragraph.runs:
        return
    for run in paragraph.runs:
        _set_run_font(run)


def _append_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(f"- {text}")
    _set_paragraph_font(p)


def _append_plain(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    _set_paragraph_font(p)


def _remove_existing_release_section(doc: Document) -> None:
    start_idx = -1
    for idx, p in enumerate(doc.paragraphs):
        if p.text.strip() == TITLE:
            start_idx = idx
            break
    if start_idx < 0:
        return

    body = doc.element.body
    for i in range(len(doc.paragraphs) - 1, start_idx - 1, -1):
        body.remove(doc.paragraphs[i]._element)


def main() -> None:
    doc = Document(DOC_PATH)
    _remove_existing_release_section(doc)

    h1 = doc.add_heading(TITLE, level=1)
    _set_paragraph_font(h1)

    h2 = doc.add_heading("14.1 自動產生報告勾稽原則（Word）", level=2)
    _set_paragraph_font(h2)
    _append_bullet(doc, "勾稽來源以照片標註（Evidence annotations.tags）為主，EXIF 不作為報告勾選依據。")
    _append_bullet(doc, "Table 1 採『先勾破壞模式，再勾致災原因』gate 原則，避免跨群誤勾。")
    _append_bullet(doc, "道路邊坡類：可勾道路上方邊坡滑動、道路下方邊坡滑動、整體性破壞，並聯動土質鬆軟、坡度過大。")
    _append_bullet(doc, "護岸/擋土牆類：已建立護岸、河道、上/下方擋土牆破壞對應規則，並聯動水文/排水/介面等致災原因。")
    _append_bullet(doc, "橋梁類：已建立橋墩、橋面、橋台、主梁等破壞模式與洪水沖刷/撞擊原因勾稽。")
    _append_bullet(doc, "Table 2 目前自動勾稽 Row1~Row5；Row6（熱危害）與 Row7（低溫危害）依現行決議維持手動不自動勾選。")
    _append_bullet(doc, "保留人工判斷欄位：Table 1 之『其他(請敘述)』與『致災原因需另辦理整體安全評估』。")
    _append_bullet(doc, "正式規格檔：docs/word_checkbox_mapping_spec.md，後續功能報告須納入本檔內容。")

    h3 = doc.add_heading("14.2 LINE 與審核流程功能（定版）", level=2)
    _set_paragraph_font(h3)
    _append_bullet(doc, "全域選單 postback 於任何 flow 下均可正確路由，不再誤跳『請選擇審核類別』。")
    _append_bullet(doc, "查詢案件 -> 選狀態 -> 查看，已可正常開啟案件詳情。")
    _append_bullet(doc, "管理者於案件詳情可執行『通過 / 退回 / 結案』；一般使用者維持權限控管。")
    _append_bullet(doc, "標註互動優化：排除選項視為單選，點擊即自動前進到下一標註分類。")
    _append_bullet(doc, "多選標註採累積模式，完成所有選項後再按一次『確認選擇』進下一步，減少重複點擊。")
    _append_bullet(doc, "災害原因選單已加入『地震』並防止重複顯示。")
    _append_bullet(doc, "必要照片順序已統一為 P1 -> P2 -> P3 -> P4，並同步調整名稱與提示文案。")

    h4 = doc.add_heading("14.3 WebGIS 與統計功能（定版）", level=2)
    _set_paragraph_font(h4)
    _append_bullet(doc, "WebGIS 圖面可視狀態統一包含：待審核、處理中、已結案、已退回。")
    _append_bullet(doc, "已退回案件若缺座標，使用工務段中心點作為低信心 fallback marker，以確保圖面可視。")
    _append_bullet(doc, "資訊摘要『最後更新』已改為完整日期時間（年月日+時分秒）。")
    _append_bullet(doc, "📊 邊坡災害統計儀表板總數與各工務段件數已改為與 WebGIS 後台一致口徑。")
    _append_bullet(doc, "統一計數來源：僅納入可視案件狀態（pending_review / in_progress / closed / returned），排除 draft。")
    _append_bullet(doc, "WebGIS 後台刪除案件後，LINE 統計摘要與 stats 儀表板會同步更新，不再出現數量不一致。")

    h5 = doc.add_heading("14.4 文件與穩定性說明", level=2)
    _set_paragraph_font(h5)
    _append_plain(doc, "本次文件使用 python-docx 直接讀寫原始 .docx 並原地覆寫，避免非 Office 相容格式造成開啟錯誤。")
    _append_plain(doc, "若使用者端仍看到舊版內容，請先關閉檔案並重新開啟或清除暫存後再讀取最新文件。")

    temp_path = DOC_PATH.with_name(DOC_PATH.stem + "_tmp.docx")
    doc.save(temp_path)
    temp_path.replace(DOC_PATH)


if __name__ == "__main__":
    main()
